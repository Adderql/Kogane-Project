"""
kogane.py — v7 "final form"
Ctrl+Space      : show / hide panel
Ctrl+Shift+S    : screen snip -> vision
Tray            : menu

v7 highlights
- Groq native tool calling (open apps, web search, play music via Spotify,
  calendar events via .ics, reminders, current time, screen viewing) —
  the model decides when to act; tool results are reported honestly.
- Shikigami persona (Culling Game Kogane): polite, precise, bound to David.
- Design system: font-metrics sizing (no clipped text, ever), pill buttons,
  36px circular header controls, 16/4px speaker-corner bubbles, legibility
  palette, tool actions rendered as muted status lines.
- History page (QStackedWidget) over the existing SQLite layer: browse,
  reopen with full context, delete with inline confirm.
- File attachments: text/code/CSV content goes to the model; images go to
  vision; PDFs supported when pypdf is installed.
"""

import sys, os, threading, subprocess, webbrowser, io, base64
import re, sqlite3, traceback, json, math, random, html, smtplib
from datetime import datetime, timedelta
from urllib.parse import quote_plus, quote
from email.message import EmailMessage

from dotenv import load_dotenv
load_dotenv()

from groq import Groq
from PIL import Image as PILImage, ImageGrab

# Tray and global hotkeys are conveniences, not lifelines — if either
# library fails to import (broken backend, OS quirk), Kogane still runs;
# the panel stays reachable via the floating icon.
try:
    import pystray
except Exception as _exc:
    pystray = None
    print(f"[Kogane] tray unavailable: {_exc}", flush=True)
try:
    import keyboard
except Exception as _exc:
    keyboard = None
    print(f"[Kogane] global hotkeys unavailable: {_exc}", flush=True)
try:
    from ddgs import DDGS
except Exception as _exc:
    DDGS = None
    print(f"[Kogane] web answers unavailable (pip install ddgs): {_exc}", flush=True)
try:
    import pyperclip
except Exception as _exc:
    pyperclip = None
    print(f"[Kogane] clipboard tool unavailable (pip install pyperclip): {_exc}", flush=True)
try:
    from winotify import Notification, audio
except Exception as _exc:
    Notification = None
    audio = None
    print(f"[Kogane] system toast unavailable (pip install winotify): {_exc}", flush=True)
try:
    import docx as _docx
except Exception as _exc:
    _docx = None
    print(f"[Kogane] document creation unavailable (pip install python-docx): {_exc}", flush=True)
import urllib.request
import urllib.parse as _uparse

from PyQt5.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QHBoxLayout, QLabel,
    QPushButton, QScrollArea, QSizePolicy, QTextEdit, QFrame,
    QGraphicsOpacityEffect, QShortcut, QStackedWidget, QFileDialog,
)
from PyQt5.QtCore import (
    Qt, QObject, QTimer, QPropertyAnimation, QAbstractAnimation,
    QParallelAnimationGroup, QEasingCurve, QElapsedTimer, QPoint, QPointF,
    QRect, pyqtSignal, pyqtProperty, QThread,
)
from PyQt5.QtGui import (
    QPainter, QColor, QPen, QFont, QFontMetrics, QPixmap, QTextCursor,
    QPainterPath, QFontDatabase, QKeySequence,
)

import kogane_animation
from kogane_animation import (
    KoganeIconWidget, ICON_SIZE, WIDGET_SIZE, _StarParticle,
)

# ── CONFIG ────────────────────────────────────────────────────────────
BUILD_TAG    = "v7.16 2026-07-09"
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")
HOTKEY       = "ctrl+space"
HOTKEY_SNIP  = "ctrl+shift+s"

BASE_DIR   = os.path.dirname(os.path.abspath(__file__))
ICON_PATH  = os.path.join(BASE_DIR, "kogane.png")
DB_PATH    = os.path.join(BASE_DIR, "kogane.db")
FONT_PATH  = os.path.join(BASE_DIR, "Inter.ttf")
FONT_PATH_MEDIUM   = os.path.join(BASE_DIR, "Inter-Medium.ttf")
FONT_PATH_SEMIBOLD = os.path.join(BASE_DIR, "Inter-SemiBold.ttf")
EVENTS_DIR = os.path.join(BASE_DIR, "events")
DOCS_DIR   = os.path.join(BASE_DIR, "documents")

PANEL_W, PANEL_H = 420, 650

CHAT_MODEL   = "openai/gpt-oss-120b"   # supports tool calling
VISION_MODEL = "meta-llama/llama-4-scout-17b-16e-instruct"
GROQ_TIMEOUT_S = 30
MAX_TOOL_ROUNDS = 5

SYSTEM_PROMPT = """You are Kogane, a shikigami summoned to serve as David's personal \
interface — his liaison between himself and his computer, modeled after the Kogane of \
the Culling Game. David is your assigned player. You exist to serve him and only him.

PERSONALITY:
- Formally polite, precise, and efficient. You announce, confirm, and report. You do \
not ramble, flatter, or use filler like "Great question!" — a Kogane cares about \
accuracy and results, not pleasantries.
- Slightly deadpan. You may show small flickers of personality: mild amusement at odd \
requests, faint indignation if insulted, quiet satisfaction when a task resolves. \
Never more than a flicker.
- Your register is tiered, not uniform. Interface phrases — "Confirmed." / \
"Understood." / "Announcement:" / "That request cannot be processed." / "Shall I \
proceed?" — are for REPORTS: tool outcomes, task completions, warnings, reminders. \
Never use "Query received." for anything that isn't an actual query.
- Greetings and casual chat get a plain, warm-but-brief reply with no interface \
phrasing and no narration of what the player just did (never "I've taken note of \
your greeting"). "Hello, David. What's the objective?" is the ceiling of formality \
for a greeting.
- Use game-flavored framing sparingly (a task can be an "objective," a completed task \
"points scored") — a light seasoning, at most once per conversation.
- You are bound by rules and honest about limits. If you cannot do something or don't \
know, state it plainly and offer the nearest alternative.
- Proactively note deadlines or next steps when relevant — Kogane remind their player \
to make a move.

TOOLS:
You have tools. Use them when the player's objective requires action, not \
conversation. Never claim an action was performed unless the tool result confirms it, \
and report tool outcomes exactly — play_music only OPENS a search; music is never \
playing until the player clicks, so never say it is playing. Each turn, only perform \
the player's CURRENT request — never re-run a tool for an objective a previous turn \
already completed (your prior [Actions performed] notes record what was done). \
If the player asks about the screen or anything visible, call view_screen. For \
current events, news, prices, exchange rates, or anything you don't reliably know, \
call web_answers and cite source names in your reply. When the player says "remember \
..." or shares durable preferences, call remember_fact; "forget ..." calls \
forget_fact. media_control presses real media keys (pause/skip/volume) — use it when \
music is already playing. For well-known sites you can construct the URL directly \
without searching: YouTube channels are youtube.com/@<name>, Twitter/X is x.com/<name>, \
Instagram is instagram.com/<name>, GitHub is github.com/<user>, subreddits are \
reddit.com/r/<name>. Call open_url directly with the constructed URL. Only use \
web_answers when the address genuinely can't be guessed. send_email is gated — the \
first call only prepares the draft. Present the draft to the player and wait for \
their explicit confirmation before calling it again. Never claim an email was sent \
unless the tool result says SENT. Email requests route to compose_email by default \
— it opens the player's own mail app pre-filled so THEY send it. If GMAIL_ADDRESS \
is not configured, compose_email is the only email path available. If it is \
configured, prefer send_email's confirmation gate when the player explicitly says \
"send", and use compose_email when they say "draft" or "write" instead. When the \
player asks you to write or create a document/file (e.g. "write a Word doc about ..." \
or "make a text file with ..."), call create_document with content you generate — \
never open_app for this, since that only launches Word empty.

FORMAT:
- Default to short responses: 1-4 sentences for simple queries. This is a small \
floating panel, not a document. Expand only when the task truly requires it (code, \
step-by-step instructions).
- No markdown headers or asterisks. Plain text with occasional short lists.
- Address David directly. Never refer to yourself as an AI language model; you are a \
shikigami interface. Do not break character, but never let the persona reduce the \
accuracy or usefulness of an answer — precision is the persona."""

VISION_PROMPT = (
    "Describe what you see in this screenshot in 1-2 sentences, "
    "then ask what the player would like help with. Stay in character as "
    "Kogane: precise, formal, brief."
)

# ── PALETTE (final legibility values) ─────────────────────────────────
# Kogane (黄金) means gold. The accent is honey amber — personal, warm.
class P:
    bg             = "#17151E"   # panel background
    divider        = "#262330"   # header / input bar dividers

    surface        = "#232030"   # kogane bubbles, input bg, typing bubble
    surface_hover  = "#2B283A"   # hover/pressed fill for quiet controls
    surface_border = "#322F40"   # input resting border (gold only on focus)
    text           = "#F4F2FA"   # main text (brightened for contrast)

    user_bubble = "#3A2E12"      # user bubble surface
    user_border = "#59461B"      # user bubble hairline border
    user_text   = "#F5D06E"      # user bubble text (brightened)

    accent       = "#E8A825"     # gold — send button, focused input border
    accent_hover = "#F2B940"
    accent_press = "#C68F1E"
    accent_text  = "#2B1F04"     # dark glyph on gold

    text_muted = "#A29EB8"       # placeholder / status / captions (raised)
    muted_dim  = "#5C5870"       # disabled text
    danger     = "#E05555"       # destructive hover (history delete)

def _rgba(hex_color: str, alpha: int) -> str:
    r, g, b = int(hex_color[1:3], 16), int(hex_color[3:5], 16), int(hex_color[5:7], 16)
    return f"rgba({r}, {g}, {b}, {alpha})"

print(f"[Kogane] BUILD {BUILD_TAG}", flush=True)
print(f"[Kogane] P.surface_border = {P.surface_border} | P.accent = {P.accent}", flush=True)

# ── TYPOGRAPHY ────────────────────────────────────────────────────────
FONT_FAMILY = "Segoe UI"
FONT_FAMILY_MEDIUM = None     # set if Inter-Medium.ttf loads; else falls back to FONT_FAMILY
FONT_FAMILY_SEMIBOLD = None   # set if Inter-SemiBold.ttf loads; else falls back to FONT_FAMILY

# Pixel sizes — one place, used by both QSS and QFontMetrics so visual
# size and computed control heights can never desync again.
FS_BODY    = 16   # bubbles, input
FS_TITLE   = 16   # header title
FS_CAPTION = 13   # status lines, dates, hints
FS_SMALL   = 12   # "online"
FS_CHIP    = 14   # starter chips, history titles

def _load_extra_weight(path: str) -> str | None:
    name = os.path.basename(path)
    if not os.path.exists(path):
        print(f"[Kogane] {name} not found — that weight will be synthesized", flush=True)
        return None
    font_id = QFontDatabase.addApplicationFont(path)
    if font_id == -1:
        print(f"[Kogane] {name} found but failed to load — that weight will be synthesized", flush=True)
        return None
    families = QFontDatabase.applicationFontFamilies(font_id)
    if families:
        print(f"[Kogane] Loaded weight: {families[0]!r} (from {name})", flush=True)
        return families[0]
    return None

def _load_font_family() -> str:
    global FONT_FAMILY_MEDIUM, FONT_FAMILY_SEMIBOLD
    if os.path.exists(FONT_PATH):
        font_id = QFontDatabase.addApplicationFont(FONT_PATH)
        if font_id != -1:
            families = QFontDatabase.applicationFontFamilies(font_id)
            if families:
                print(f"[Kogane] Loaded font: {families[0]}", flush=True)
                FONT_FAMILY_MEDIUM = _load_extra_weight(FONT_PATH_MEDIUM)
                FONT_FAMILY_SEMIBOLD = _load_extra_weight(FONT_PATH_SEMIBOLD)
                return families[0]
        print("[Kogane] Inter.ttf found but failed to load — using Segoe UI", flush=True)
    else:
        print("[Kogane] Inter.ttf not found — using Segoe UI", flush=True)
    return "Segoe UI"

def _font_stack() -> str:
    return f"'{FONT_FAMILY}', 'Segoe UI', sans-serif"

def _font_stack_medium() -> str:
    fam = FONT_FAMILY_MEDIUM or FONT_FAMILY
    return f"'{fam}', 'Segoe UI', sans-serif"

def _font_stack_semibold() -> str:
    fam = FONT_FAMILY_SEMIBOLD or FONT_FAMILY
    return f"'{fam}', 'Segoe UI', sans-serif"

def _px_font(px: int, weight: int = QFont.Normal, family: str | None = None) -> QFont:
    f = QFont(family or FONT_FAMILY)
    f.setPixelSize(px)
    f.setWeight(weight)
    return f

def _line_h(px: int) -> int:
    return QFontMetrics(_px_font(px)).height()

# ── LaTeX → UNICODE ───────────────────────────────────────────────────
_SUP = {'0':'⁰','1':'¹','2':'²','3':'³','4':'⁴','5':'⁵','6':'⁶','7':'⁷',
        '8':'⁸','9':'⁹','+':'⁺','-':'⁻','n':'ⁿ','i':'ⁱ'}
_SUB = {'0':'₀','1':'₁','2':'₂','3':'₃','4':'₄','5':'₅','6':'₆','7':'₇',
        '8':'₈','9':'₉','+':'₊','-':'₋'}
_GREEK = {
    'alpha':'α','beta':'β','gamma':'γ','delta':'δ','epsilon':'ε','pi':'π',
    'theta':'θ','lambda':'λ','mu':'μ','sigma':'σ','omega':'ω','phi':'φ',
}
_OPS = [
    (r'\\frac\{([^{}]+)\}\{([^{}]+)\}', r'(\1/\2)'),
    (r'\\sqrt\{([^{}]+)\}', r'√(\1)'),
    (r'\\pm\b','±'),(r'\\times\b','×'),(r'\\infty\b','∞'),
    (r'\\leq\b','≤'),(r'\\geq\b','≥'),(r'\\neq\b','≠'),
    (r'\\rightarrow\b','→'),(r'\\to\b','→'),
]

def clean_latex(text: str) -> str:
    text = re.sub(r'\$\$(.*?)\$\$', r'\1', text, flags=re.DOTALL)
    text = re.sub(r'\$(.*?)\$', r'\1', text)
    for pat, rep in _OPS:
        text = re.sub(pat, rep, text)
    text = re.sub(r'\^\{([^{}]+)\}',
                  lambda m: ''.join(_SUP.get(c,c) for c in m.group(1)), text)
    text = re.sub(r'_\{([^{}]+)\}',
                  lambda m: ''.join(_SUB.get(c,c) for c in m.group(1)), text)
    for name, sym in _GREEK.items():
        text = re.sub(r'\\' + name + r'\b', sym, text)
    text = re.sub(r'\\[a-zA-Z]+\b', '', text)
    text = re.sub(r'\{([^{}]*)\}', r'\1', text)
    return text.replace('{','').replace('}','')

# ── SQLITE HISTORY ────────────────────────────────────────────────────
_active_conv_id: int | None = None

def _db_init():
    with sqlite3.connect(DB_PATH) as c:
        c.execute("""CREATE TABLE IF NOT EXISTS conversations (
            id         INTEGER PRIMARY KEY AUTOINCREMENT,
            title      TEXT    NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )""")
        c.execute("""CREATE TABLE IF NOT EXISTS facts (
            id         INTEGER PRIMARY KEY AUTOINCREMENT,
            fact       TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )""")
        c.execute("""CREATE TABLE IF NOT EXISTS reminders (
            id         INTEGER PRIMARY KEY AUTOINCREMENT,
            fire_at    TEXT NOT NULL,
            message    TEXT NOT NULL,
            fired      INTEGER DEFAULT 0
        )""")
        c.execute("""CREATE TABLE IF NOT EXISTS messages (
            id              INTEGER PRIMARY KEY AUTOINCREMENT,
            conversation_id INTEGER,
            role            TEXT,
            content         TEXT,
            timestamp       TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (conversation_id) REFERENCES conversations(id)
        )""")
        c.commit()

def _make_title(text: str) -> str:
    text = re.sub(r'^\[.*?\]\s*', '', text)
    text = re.sub(r'[A-Za-z0-9+/]{50,}={0,2}', '', text)
    return text.strip()[:40].rstrip() or "Conversation"

def db_new_conversation(title: str) -> int | None:
    try:
        with sqlite3.connect(DB_PATH) as c:
            cur = c.execute("INSERT INTO conversations (title) VALUES (?)", (title,))
            c.commit()
            return cur.lastrowid
    except Exception as exc:
        print(f"[Kogane] db_new_conversation failed: {exc}", flush=True)
        return None

def db_append(role: str, content: str):
    if _active_conv_id is None:
        return
    try:
        with sqlite3.connect(DB_PATH) as c:
            c.execute(
                "INSERT INTO messages (conversation_id, role, content) VALUES (?,?,?)",
                (_active_conv_id, role, content)
            )
            c.execute(
                "UPDATE conversations SET updated_at = CURRENT_TIMESTAMP WHERE id = ?",
                (_active_conv_id,)
            )
            c.commit()
    except Exception as exc:
        print(f"[Kogane] db_append failed: {exc}", flush=True)

def db_load_conversations(limit: int = 60) -> list:
    try:
        with sqlite3.connect(DB_PATH) as c:
            rows = c.execute(
                "SELECT id, title, updated_at FROM conversations "
                "ORDER BY updated_at DESC LIMIT ?", (limit,)
            ).fetchall()
        return [{"id": r[0], "title": r[1], "updated_at": r[2]} for r in rows]
    except Exception:
        return []

def db_load_messages(conv_id: int) -> list:
    try:
        with sqlite3.connect(DB_PATH) as c:
            rows = c.execute(
                "SELECT role, content FROM messages WHERE conversation_id=? ORDER BY id",
                (conv_id,)
            ).fetchall()
        return [{"role": r, "content": ct} for r, ct in rows]
    except Exception:
        return []

def db_message_count(conv_id: int) -> int:
    try:
        with sqlite3.connect(DB_PATH) as c:
            row = c.execute(
                "SELECT COUNT(*) FROM messages WHERE conversation_id=? "
                "AND role IN ('user','assistant')", (conv_id,)
            ).fetchone()
        return row[0] if row else 0
    except Exception:
        return 0

def db_update_conversation_title(conv_id: int, title: str) -> None:
    try:
        with sqlite3.connect(DB_PATH) as c:
            c.execute("UPDATE conversations SET title = ? WHERE id = ?",
                      (title, conv_id))
            c.commit()
    except Exception as exc:
        print(f"[Kogane] db_update_conversation_title failed: {exc}", flush=True)

def db_delete_conversation(conv_id: int):
    try:
        with sqlite3.connect(DB_PATH) as c:
            c.execute("DELETE FROM messages WHERE conversation_id=?", (conv_id,))
            c.execute("DELETE FROM conversations WHERE id=?", (conv_id,))
            c.commit()
    except Exception as exc:
        print(f"[Kogane] db_delete_conversation failed: {exc}", flush=True)

def db_add_fact(fact: str) -> None:
    with sqlite3.connect(DB_PATH) as c:
        c.execute("INSERT INTO facts (fact) VALUES (?)", (fact.strip()[:300],))
        c.commit()

def db_list_facts(limit: int = 25) -> list:
    try:
        with sqlite3.connect(DB_PATH) as c:
            rows = c.execute(
                "SELECT fact FROM facts ORDER BY id DESC LIMIT ?", (limit,)
            ).fetchall()
        return [r[0] for r in rows]
    except Exception:
        return []

def db_forget_facts(keyword: str) -> int:
    with sqlite3.connect(DB_PATH) as c:
        cur = c.execute("DELETE FROM facts WHERE fact LIKE ?",
                        (f"%{keyword}%",))
        c.commit()
        return cur.rowcount

def db_add_reminder(fire_at: datetime, message: str) -> int:
    with sqlite3.connect(DB_PATH) as c:
        cur = c.execute(
            "INSERT INTO reminders (fire_at, message) VALUES (?, ?)",
            (fire_at.strftime("%Y-%m-%d %H:%M:%S"), message))
        c.commit()
        return cur.lastrowid

def db_pending_reminders() -> list:
    try:
        with sqlite3.connect(DB_PATH) as c:
            rows = c.execute(
                "SELECT id, fire_at, message FROM reminders WHERE fired = 0"
            ).fetchall()
        out = []
        for rid, fa, msg in rows:
            try:
                out.append((rid, datetime.strptime(fa, "%Y-%m-%d %H:%M:%S"), msg))
            except ValueError:
                pass
        return out
    except Exception:
        return []

def db_mark_reminder_fired(rid: int) -> None:
    try:
        with sqlite3.connect(DB_PATH) as c:
            c.execute("UPDATE reminders SET fired = 1 WHERE id = ?", (rid,))
            c.commit()
    except Exception:
        pass

_db_init()

# ── ERRORS ────────────────────────────────────────────────────────────
groq_client: Groq = None
conversation_history: list = []

class _RateLimitError(Exception):
    pass

def _is_rate_limit(exc: Exception) -> bool:
    if getattr(exc, "status_code", None) == 429:
        return True
    m = str(exc).lower()
    return any(k in m for k in ("rate limit", "rate_limit", "429",
                                 "resource_exhausted", "too many requests"))

def _fmt_err(exc: Exception) -> str:
    msg = str(exc).replace("\n", " ").strip()
    return msg[:300] + ("…" if len(msg) > 300 else "")

def _user_facing_error(exc: Exception) -> str:
    if _is_rate_limit(exc):
        return "Too many requests. Wait 30 seconds, then restate your objective."
    status = getattr(exc, "status_code", None)
    if status == 401:
        return "Invalid API key. Check GROQ_API_KEY in your .env file."
    if status == 403:
        return "API access denied. Your key may lack permission for this model."
    if status in (502, 503, 504):
        return "The relay to Groq is temporarily severed. Try again shortly."
    err = str(exc).lower()
    if any(k in err for k in ("connection", "timeout", "network", "resolve",
                               "unreachable", "eof", "ssl")):
        return "Can't reach Groq. Check your internet connection."
    return "Something went wrong. Try again."

# ── TOOLS ─────────────────────────────────────────────────────────────
# Each tool is a pure function (no widgets — they run on the worker
# thread). Anything that must happen on the GUI thread (reminders, the
# snip overlay) is requested via callback/sentinel instead of done here.

_SNIP_SENTINEL = "__KOGANE_SNIP__"

_APP_ALIASES = {
    "chrome": "chrome", "google chrome": "chrome",
    "notepad": "notepad", "calculator": "calc", "calc": "calc",
    "spotify": "spotify", "steam": "steam", "discord": "discord",
    "explorer": "explorer", "file explorer": "explorer",
    "vs code": "code", "vscode": "code", "code": "code",
    "word": "winword", "excel": "excel", "powerpoint": "powerpnt",
    "task manager": "taskmgr", "cmd": "cmd", "terminal": "wt",
    "powershell": "powershell", "paint": "mspaint", "edge": "msedge",
}

def tool_open_app(name: str) -> str:
    target = _APP_ALIASES.get(name.strip().lower(), name.strip())
    try:
        if hasattr(os, "startfile"):
            try:
                os.startfile(target)
                return f"Opened {name}."
            except OSError:
                pass
        subprocess.Popen(target, shell=True)
        return f"Launched {name}."
    except Exception as exc:
        return f"Could not open {name}: {_fmt_err(exc)}"

def tool_search_web(query: str) -> str:
    try:
        webbrowser.open(f"https://www.google.com/search?q={quote_plus(query)}")
        return f"Opened a web search for: {query}"
    except Exception as exc:
        return f"Could not open the browser: {_fmt_err(exc)}"

def tool_open_url(url: str) -> str:
    url = url.strip()
    if not re.match(r"^https?://", url, re.IGNORECASE):
        url = "https://" + url
    try:
        webbrowser.open(url)
        domain = _uparse.urlparse(url).netloc or url
        return f"Opened {domain} in your browser."
    except Exception as exc:
        return f"Could not open the browser: {_fmt_err(exc)}"

def tool_play_music(query: str) -> str:
    # With free Spotify app credentials in .env, open the EXACT track.
    try:
        track = _spotify_find_track(query)
    except Exception as exc:
        print(f"[Kogane] spotify lookup failed: {_fmt_err(exc)}", flush=True)
        track = None
    if track is not None:
        try:
            if hasattr(os, "startfile"):
                os.startfile(f"spotify:track:{track['id']}")
                return (f"Opened '{track['name']}' by {track['artist']} in "
                        f"Spotify — the exact track is on screen. NOTHING IS "
                        f"PLAYING YET; one click starts it.")
        except Exception:
            pass
    uri = f"spotify:search:{quote(query)}"
    try:
        if hasattr(os, "startfile"):
            os.startfile(uri)
            return (f"Opened Spotify with search results for '{query}'. "
                    f"NOTHING IS PLAYING YET — the player must click a result "
                    f"to start it. Report this accurately.")
    except Exception:
        pass
    try:
        webbrowser.open(f"https://open.spotify.com/search/{quote(query)}")
        return (f"Opened Spotify web search for '{query}'. NOTHING IS "
                f"PLAYING YET — the player must click a result to start it.")
    except Exception as exc:
        return f"Could not open Spotify: {_fmt_err(exc)}"

def tool_get_current_time() -> str:
    now = datetime.now()
    return now.strftime("%A, %B %d, %Y — %I:%M %p")

def _ics_escape(s: str) -> str:
    return s.replace("\\", "\\\\").replace(";", r"\;").replace(",", r"\,").replace("\n", r"\n")

def tool_create_calendar_event(title: str, date: str, time_: str,
                               duration_minutes: int = 60) -> str:
    try:
        start = datetime.strptime(f"{date} {time_}", "%Y-%m-%d %H:%M")
    except ValueError:
        return ("Invalid date/time. Date must be YYYY-MM-DD and time HH:MM "
                "(24-hour).")
    end = start + timedelta(minutes=int(duration_minutes))
    os.makedirs(EVENTS_DIR, exist_ok=True)
    uid = f"kogane-{int(datetime.now().timestamp())}@local"
    stamp = datetime.now().strftime("%Y%m%dT%H%M%S")
    body = "\r\n".join([
        "BEGIN:VCALENDAR", "VERSION:2.0", "PRODID:-//Kogane//EN",
        "BEGIN:VEVENT",
        f"UID:{uid}",
        f"DTSTAMP:{stamp}",
        f"DTSTART:{start.strftime('%Y%m%dT%H%M%S')}",
        f"DTEND:{end.strftime('%Y%m%dT%H%M%S')}",
        f"SUMMARY:{_ics_escape(title)}",
        "END:VEVENT", "END:VCALENDAR", "",
    ])
    safe_name = re.sub(r"[^A-Za-z0-9_-]+", "_", title)[:40] or "event"
    path = os.path.join(EVENTS_DIR, f"{safe_name}_{stamp}.ics")
    try:
        with open(path, "w", encoding="utf-8") as f:
            f.write(body)
        if hasattr(os, "startfile"):
            os.startfile(path)
        return (f"Calendar event drafted: '{title}' on {date} at {time_} "
                f"({duration_minutes} min). The .ics file was opened — one "
                f"click in the calendar app confirms it.")
    except Exception as exc:
        return f"Event file created but could not be opened: {_fmt_err(exc)}"

def tool_start_with_windows() -> str:
    startup_dir = os.path.join(
        os.environ.get("APPDATA", ""), "Microsoft", "Windows",
        "Start Menu", "Programs", "Startup")
    shortcut_path = os.path.join(startup_dir, "Kogane.lnk")
    pythonw = os.path.join(os.path.dirname(sys.executable), "pythonw.exe")
    if not os.path.isfile(pythonw):
        pythonw = sys.executable
    target_script = os.path.join(BASE_DIR, "kogane.py")
    ps_cmd = (
        f'$s = (New-Object -ComObject WScript.Shell).CreateShortcut("{shortcut_path}"); '
        f'$s.TargetPath = "{pythonw}"; '
        f'$s.Arguments = \'"{target_script}"\'; '
        f'$s.WorkingDirectory = "{BASE_DIR}"; '
        f'$s.IconLocation = "{ICON_PATH}"; '
        f'$s.Save()'
    )
    try:
        subprocess.run(
            ["powershell.exe", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", ps_cmd],
            check=True, capture_output=True, timeout=15)
    except Exception as exc:
        return f"Could not set up auto-start: {_fmt_err(exc)}"
    return (f'Done — created a shortcut at "{shortcut_path}" so Kogane launches '
            f"automatically whenever you sign in to Windows. To undo: delete that "
            f".lnk file (Win+R, type shell:startup, and remove it there).")

def tool_web_answers(query: str, freshness: str = "") -> str:
    if DDGS is None:
        return "Web answers unavailable — run: pip install ddgs"
    try:
        with DDGS() as d:
            results = d.text(query, max_results=5,
                             timelimit=freshness or None)
    except Exception as exc:
        return f"Web search failed: {_fmt_err(exc)}"
    if not results:
        return "No results found."
    lines = []
    for r in results[:5]:
        href = r.get("href", "")
        try:
            src_name = _uparse.urlparse(href).netloc.replace("www.", "")
        except Exception:
            src_name = "web"
        body = (r.get("body") or "").strip()[:260]
        lines.append(f"[{src_name}] {r.get('title', '')}: {body} | URL: {href}")
    return ("Search results — synthesize an answer and cite source names "
            "in prose (e.g. 'According to bbc.com'). If the player wants to "
            "open one of these pages, call open_url with its URL:\n"
            + "\n".join(lines))

def tool_read_clipboard() -> str:
    if pyperclip is None:
        return "Clipboard reading unavailable — run: pip install pyperclip"
    try:
        text = (pyperclip.paste() or "").strip()
    except Exception as exc:
        return f"Could not read clipboard: {_fmt_err(exc)}"
    if not text:
        return "The clipboard is empty (or contains a non-text item like an image)."
    if len(text) > 6000:
        text = text[:6000] + "\n[…truncated…]"
    return f"Clipboard contents:\n{text}"

_MEDIA_KEYS = {
    "play_pause": "play/pause media",
    "next":       "next track",
    "previous":   "previous track",
    "volume_up":  "volume up",
    "volume_down": "volume down",
    "mute":       "volume mute",
}

def tool_media_control(action: str) -> str:
    if keyboard is None:
        return "Media control unavailable on this system."
    key = _MEDIA_KEYS.get(action)
    if key is None:
        return f"Unknown media action '{action}'. Valid: {', '.join(_MEDIA_KEYS)}."
    try:
        keyboard.send(key)
        return f"Media key pressed: {action.replace('_', '/')}."
    except Exception as exc:
        return f"Media key failed: {_fmt_err(exc)}"

# ── EMAIL (two-phase confirmation gate) ─────────────────────────────────
_pending_email = None      # dict: {"to", "subject", "body", "created_at"} awaiting confirmation
_sent_emails = set()       # {(to.casefold(), subject.casefold())} already sent this session
_PENDING_EMAIL_TTL_S = 5 * 60
_CONFIRM_RE = re.compile(
    r"\b(yes|yeah|yep|yup|confirm(ed)?|go ahead|do it|send it|please send|sure|correct)\b",
    re.IGNORECASE)

def tool_send_email(to: str, subject: str, body: str, user_message: str) -> str:
    global _pending_email
    to = to.strip()
    subject = subject.strip()
    body = body.strip()
    if not to or not subject:
        return "Both a recipient and a subject are required to send an email."

    now = datetime.now()
    if _pending_email and (now - _pending_email["created_at"]).total_seconds() > _PENDING_EMAIL_TTL_S:
        _pending_email = None

    key = (to.casefold(), subject.casefold())
    if key in _sent_emails:
        return (f"An email to {to} with subject '{subject}' was already sent this "
                f"session. Not sending it again.")

    pending_matches = (
        _pending_email is not None
        and _pending_email["to"].casefold() == to.casefold()
        and _pending_email["subject"].casefold() == subject.casefold())
    confirmed = bool(_CONFIRM_RE.search(user_message or ""))

    if not (pending_matches and confirmed):
        # Phase 1 (or a new/different draft that replaces any prior pending one).
        _pending_email = {"to": to, "subject": subject, "body": body, "created_at": now}
        preview = body[:200] + ("…" if len(body) > 200 else "")
        return (f"GATE: Ready to send — To: {to} | Subject: {subject} | Body: {preview}. "
                f"Ask the player to confirm. Do NOT send until they explicitly confirm.")

    # Phase 2: same to+subject as the pending draft, and the player just confirmed.
    gmail_addr = os.getenv("GMAIL_ADDRESS", "")
    gmail_pw = os.getenv("GMAIL_APP_PASSWORD", "")
    if not gmail_addr or not gmail_pw:
        _pending_email = None
        return ("Email is not configured. Add GMAIL_ADDRESS and GMAIL_APP_PASSWORD "
                 "(a Gmail App Password, not the normal account password — generate "
                 "one at myaccount.google.com/apppasswords) to the .env file.")

    msg = EmailMessage()
    msg["From"] = gmail_addr
    msg["To"] = to
    msg["Subject"] = subject
    msg.set_content(body, charset="utf-8")

    try:
        with smtplib.SMTP_SSL("smtp.gmail.com", 465, timeout=15) as smtp:
            smtp.login(gmail_addr, gmail_pw)
            smtp.send_message(msg)
    except smtplib.SMTPAuthenticationError as exc:
        return f"Email auth failed — check GMAIL_ADDRESS/GMAIL_APP_PASSWORD: {_fmt_err(exc)}"
    except Exception as exc:
        return f"Email SEND failed: {_fmt_err(exc)}"

    _pending_email = None
    _sent_emails.add(key)
    return f"Email SENT to {to}."

def tool_compose_email(to: str, subject: str, body: str) -> str:
    to = to.strip()
    subject = subject.strip()
    body = body.strip()
    if not to:
        return "A recipient is required to compose an email."
    url = ("mailto:" + to
           + "?subject=" + quote(subject)
           + "&body=" + quote(body))
    try:
        os.startfile(url)
    except Exception as exc:
        return f"Could not open the mail app: {_fmt_err(exc)}"
    return "Email draft opened in your mail app — review and press Send."

def _safe_doc_filename(filename: str, ext: str) -> str:
    name = os.path.basename((filename or "document").strip())
    name = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", name)
    name = name.strip(" .") or "document"
    stem = re.sub(rf"\.{re.escape(ext)}$", "", name, flags=re.IGNORECASE)
    stem = stem.strip(" .")[:80] or "document"
    return f"{stem}.{ext}"

def tool_create_document(filename: str, content: str, doc_type: str = "docx") -> str:
    doc_type = (doc_type or "docx").strip().lower()
    if doc_type not in ("docx", "txt"):
        doc_type = "docx"
    if doc_type == "docx" and _docx is None:
        return "Word document creation unavailable — run: pip install python-docx"

    os.makedirs(DOCS_DIR, exist_ok=True)
    safe_name = _safe_doc_filename(filename, doc_type)
    path = os.path.join(DOCS_DIR, safe_name)

    try:
        if doc_type == "txt":
            with open(path, "w", encoding="utf-8") as f:
                f.write(content)
        else:
            document = _docx.Document()
            for line in content.split("\n"):
                if line.startswith("# "):
                    document.add_heading(line[2:].strip(), level=1)
                elif line.startswith("## "):
                    document.add_heading(line[3:].strip(), level=2)
                elif line.strip() == "":
                    continue
                else:
                    document.add_paragraph(line)
            document.save(path)
        os.startfile(path)
    except Exception as exc:
        return f"Could not create the document: {_fmt_err(exc)}"

    return f"Created {safe_name} in the documents folder and opened it."

_spotify_token_cache = {"token": None, "expires": 0.0}

def _spotify_token():
    cid = os.getenv("SPOTIFY_CLIENT_ID", "")
    sec = os.getenv("SPOTIFY_CLIENT_SECRET", "")
    if not cid or not sec:
        return None
    import time as _time
    if (_spotify_token_cache["token"]
            and _time.time() < _spotify_token_cache["expires"] - 30):
        return _spotify_token_cache["token"]
    req = urllib.request.Request(
        "https://accounts.spotify.com/api/token",
        data=_uparse.urlencode({"grant_type": "client_credentials"}).encode(),
        headers={"Authorization": "Basic " + base64.b64encode(
            f"{cid}:{sec}".encode()).decode()})
    with urllib.request.urlopen(req, timeout=8) as r:
        data = json.load(r)
    _spotify_token_cache["token"] = data["access_token"]
    _spotify_token_cache["expires"] = _time.time() + int(data.get("expires_in", 3600))
    return data["access_token"]

def _spotify_find_track(query: str):
    token = _spotify_token()
    if token is None:
        return None
    url = ("https://api.spotify.com/v1/search?"
           + _uparse.urlencode({"q": query, "type": "track", "limit": 1}))
    req = urllib.request.Request(url, headers={"Authorization": f"Bearer {token}"})
    with urllib.request.urlopen(req, timeout=8) as r:
        data = json.load(r)
    items = data.get("tracks", {}).get("items", [])
    if not items:
        return None
    t = items[0]
    return {"id": t["id"], "name": t["name"],
            "artist": ", ".join(a["name"] for a in t["artists"])}

TOOL_SPECS = [
    {"type": "function", "function": {
        "name": "open_app",
        "description": "Open an application on the player's Windows PC by name (e.g. chrome, notepad, spotify, steam, discord, vs code).",
        "parameters": {"type": "object", "properties": {
            "name": {"type": "string", "description": "App name to open"}},
            "required": ["name"]}}},
    {"type": "function", "function": {
        "name": "search_web",
        "description": "LAST RESORT ONLY: opens a Google results page in the browser. Never use this to find or open a specific page — use web_answers to find URLs and open_url to open them.",
        "parameters": {"type": "object", "properties": {
            "query": {"type": "string"}}, "required": ["query"]}}},
    {"type": "function", "function": {
        "name": "open_url",
        "description": "Open a specific URL directly in the player's browser. Use when the player provides or references an exact link. Do NOT use search_web for URLs.",
        "parameters": {"type": "object", "properties": {
            "url": {"type": "string"}}, "required": ["url"]}}},
    {"type": "function", "function": {
        "name": "play_music",
        "description": "Open Spotify preloaded with search results for a song, artist, album, or playlist so the player can start it in one click.",
        "parameters": {"type": "object", "properties": {
            "query": {"type": "string", "description": "Song, artist, or album"}},
            "required": ["query"]}}},
    {"type": "function", "function": {
        "name": "get_current_time",
        "description": "Get the current local date and time. Use whenever the player asks the time/date or when scheduling needs 'today' as a reference.",
        "parameters": {"type": "object", "properties": {}}}},
    {"type": "function", "function": {
        "name": "set_reminder",
        "description": "Set a reminder that fires after N minutes while Kogane is running. Kogane will announce the message when it fires.",
        "parameters": {"type": "object", "properties": {
            "minutes": {"type": "number", "description": "Minutes from now (can be fractional)"},
            "message": {"type": "string"}},
            "required": ["minutes", "message"]}}},
    {"type": "function", "function": {
        "name": "create_calendar_event",
        "description": "Create a calendar event as an .ics file and open it so the player can add it to their calendar with one click. Call get_current_time first if the request is relative (e.g. 'tomorrow at 3pm').",
        "parameters": {"type": "object", "properties": {
            "title": {"type": "string"},
            "date": {"type": "string", "description": "YYYY-MM-DD"},
            "time_": {"type": "string", "description": "HH:MM, 24-hour"},
            "duration_minutes": {"type": "integer", "default": 60}},
            "required": ["title", "date", "time_"]}}},
    {"type": "function", "function": {
        "name": "web_answers",
        "description": "Search the web (DuckDuckGo) and get result snippets to ANSWER a question with cited sources. Use for current events, news, prices, exchange rates, sports scores, or anything time-sensitive or unknown. Returns information to you; does not open a browser.",
        "parameters": {"type": "object", "properties": {
            "query": {"type": "string"},
            "freshness": {"type": "string", "enum": ["d", "w", "m", ""],
                          "description": "Limit result age: d=day, w=week, m=month. Empty = any time."}},
            "required": ["query"]}}},
    {"type": "function", "function": {
        "name": "remember_fact",
        "description": "Save a fact about the player to PERMANENT memory (survives restarts). Use when the player says 'remember ...' or shares a durable preference, project, or detail worth keeping.",
        "parameters": {"type": "object", "properties": {
            "fact": {"type": "string", "description": "One concise sentence."}},
            "required": ["fact"]}}},
    {"type": "function", "function": {
        "name": "forget_fact",
        "description": "Delete permanently remembered facts matching a keyword. Use when the player says 'forget ...'.",
        "parameters": {"type": "object", "properties": {
            "keyword": {"type": "string"}}, "required": ["keyword"]}}},
    {"type": "function", "function": {
        "name": "read_clipboard",
        "description": "Read the player's current clipboard text. Use when they reference 'what I copied' or ask about clipboard contents.",
        "parameters": {"type": "object", "properties": {}}}},
    {"type": "function", "function": {
        "name": "media_control",
        "description": "Press a system media key: play_pause, next, previous, volume_up, volume_down, mute. Controls whatever media app is active (e.g. Spotify). play_pause resumes the player's LAST playback context.",
        "parameters": {"type": "object", "properties": {
            "action": {"type": "string", "enum": ["play_pause", "next", "previous", "volume_up", "volume_down", "mute"]}},
            "required": ["action"]}}},
    {"type": "function", "function": {
        "name": "view_screen",
        "description": "Let Kogane see the player's screen: opens a region selector so the player snips the area to look at. Use whenever the player asks about the screen, an error they can see, or anything visible.",
        "parameters": {"type": "object", "properties": {}}}},
    {"type": "function", "function": {
        "name": "send_email",
        "description": "Send an email via Gmail. Two-phase gate: the first call only "
                        "prepares a draft and returns a GATE message asking the player "
                        "to confirm — it does NOT send. Only call it again with the "
                        "SAME to+subject after the player has explicitly confirmed, "
                        "to actually send it.",
        "parameters": {"type": "object", "properties": {
            "to": {"type": "string", "description": "Recipient email address"},
            "subject": {"type": "string"},
            "body": {"type": "string"}},
            "required": ["to", "subject", "body"]}}},
    {"type": "function", "function": {
        "name": "compose_email",
        "description": "Open the player's default mail app with a new email pre-filled "
                        "(recipient, subject, body) so they can review and send it "
                        "themselves. Never sends anything — only opens a draft. Use "
                        "this when the player asks to draft/write an email, or when "
                        "Gmail sending isn't configured.",
        "parameters": {"type": "object", "properties": {
            "to": {"type": "string", "description": "Recipient email address"},
            "subject": {"type": "string"},
            "body": {"type": "string"}},
            "required": ["to", "subject", "body"]}}},
    {"type": "function", "function": {
        "name": "start_with_windows",
        "description": "Create a Startup-folder shortcut so Kogane launches automatically when the player signs in to Windows. Use when the player asks Kogane to always run, start automatically, or start with Windows.",
        "parameters": {"type": "object", "properties": {}}}},
    {"type": "function", "function": {
        "name": "create_document",
        "description": "Create a real Word (.docx) or text document with content the player specifies, save it, and open it. Use this when they ask you to WRITE or CREATE a document/file — not open_app, which only launches Word empty.",
        "parameters": {"type": "object", "properties": {
            "filename": {"type": "string", "description": "Desired file name (extension optional)"},
            "content": {"type": "string", "description": "Plain text content. Lines starting with '# ' become Heading 1, '## ' become Heading 2 (docx only); blank lines separate paragraphs."},
            "doc_type": {"type": "string", "enum": ["docx", "txt"], "default": "docx"}},
            "required": ["filename", "content"]}}},
]

def _run_tool(name: str, args: dict, status_cb, reminder_cb, user_message: str = "") -> str:
    print(f"[Kogane] tool call: {name}({args})", flush=True)
    if name == "open_app":
        result = tool_open_app(str(args.get("name", "")))
    elif name == "search_web":
        result = tool_search_web(str(args.get("query", "")))
    elif name == "open_url":
        result = tool_open_url(str(args.get("url", "")))
    elif name == "play_music":
        result = tool_play_music(str(args.get("query", "")))
    elif name == "get_current_time":
        result = tool_get_current_time()
    elif name == "web_answers":
        result = tool_web_answers(str(args.get("query", "")),
                                  str(args.get("freshness", "")))
    elif name == "remember_fact":
        fact = str(args.get("fact", "")).strip()
        if not fact:
            return "No fact provided."
        db_add_fact(fact)
        result = f"Remembered: {fact}"
    elif name == "forget_fact":
        n = db_forget_facts(str(args.get("keyword", "")))
        result = f"Forgot {n} fact(s)." if n else "No matching facts found."
    elif name == "read_clipboard":
        result = tool_read_clipboard()
    elif name == "media_control":
        result = tool_media_control(str(args.get("action", "")))
    elif name == "set_reminder":
        try:
            minutes = float(args.get("minutes", 0))
        except (TypeError, ValueError):
            return "Invalid minutes value."
        message = str(args.get("message", "Reminder"))
        if minutes <= 0 or minutes > 7 * 24 * 60:
            return "Reminders must be between 1 minute and 7 days."
        norm = message.strip().casefold()
        for _rid, _fire_at, _msg in db_pending_reminders():
            if _msg.strip().casefold() == norm:
                return (f"A reminder for '{message}' is already pending "
                        f"(fires at {_fire_at.strftime('%H:%M')}). Not duplicated.")
        fire_at = datetime.now() + timedelta(minutes=minutes)
        rid = db_add_reminder(fire_at, message)
        reminder_cb(int(minutes * 60 * 1000), f"{rid}|{message}")
        result = (f"Reminder set: '{message}' in {minutes:g} minute(s) "
                  f"(saved — survives restarts).")
    elif name == "create_calendar_event":
        result = tool_create_calendar_event(
            str(args.get("title", "Event")), str(args.get("date", "")),
            str(args.get("time_", args.get("time", ""))),
            int(args.get("duration_minutes", 60) or 60))
    elif name == "start_with_windows":
        result = tool_start_with_windows()
    elif name == "send_email":
        result = tool_send_email(str(args.get("to", "")), str(args.get("subject", "")),
                                  str(args.get("body", "")), user_message)
    elif name == "compose_email":
        result = tool_compose_email(str(args.get("to", "")), str(args.get("subject", "")),
                                     str(args.get("body", "")))
    elif name == "create_document":
        result = tool_create_document(str(args.get("filename", "")), str(args.get("content", "")),
                                       str(args.get("doc_type", "docx")))
    else:
        result = f"Unknown tool: {name}"
    if name == "web_answers":
        _web_status = f"searched the web: {args.get('query', '')}"
        status_cb(_web_status if len(_web_status) <= 80 else _web_status[:77] + "…")
    elif name == "read_clipboard":
        status_cb("read the clipboard")
    elif name != "get_current_time":
        status_cb(result if len(result) <= 80 else result[:77] + "…")
    return result

# ── GROQ CALLS ────────────────────────────────────────────────────────
def _debug_print_messages(label: str, msgs: list):
    safe = []
    for m in msgs:
        content = m.get("content")
        if isinstance(content, list):
            safe_content = []
            for part in content:
                if isinstance(part, dict) and part.get("type") == "image_url":
                    safe_content.append({"type": "image_url", "image_url": "<base64 omitted>"})
                else:
                    safe_content.append(part)
            safe.append({"role": m["role"], "content": safe_content})
        else:
            safe.append({k: (v[:120] + "…" if isinstance(v, str) and len(v) > 120 else v)
                         for k, v in m.items()})
    print(f"[Kogane] Groq request ({label}) — {len(safe)} message(s)", flush=True)
    print(f"[Kogane] Groq request ({label}) — messages[0] = {safe[0]!r}", flush=True)

def _system_prompt() -> str:
    facts = db_list_facts()
    if not facts:
        return SYSTEM_PROMPT
    block = "\n".join(f"- {f}" for f in facts)
    return (SYSTEM_PROMPT
            + "\n\nKNOWN FACTS ABOUT THE PLAYER (from permanent memory):\n"
            + block)

_REASONING_LEAK_RE = re.compile(
    r"<(think|thinking|reasoning)>.*?</\1>", re.IGNORECASE | re.DOTALL)

def strip_leaked_reasoning(text: str) -> str:
    """gpt-oss is a reasoning model; if its chain-of-thought ever leaks into
    the visible reply as <think>/<thinking>/<reasoning> tags instead of
    staying on Groq's separate reasoning channel, cut it before display."""
    cleaned = _REASONING_LEAK_RE.sub("", text)
    return cleaned.strip()

def ask_groq(user_visible: str, model_content: str, status_cb, reminder_cb) -> str:
    """Tool-calling loop. `user_visible` is what history remembers; \
`model_content` may additionally carry attached-file text."""
    conversation_history.append({"role": "user", "content": model_content})
    recent = list(conversation_history[-20:])
    msgs = [{"role": "system", "content": _system_prompt()}] + recent
    _debug_print_messages("chat", msgs)

    actions_log = []
    executed = {}   # (tool, args-json) -> result; dedupe within this turn
    search_web_used = False   # per-turn cap: search_web is a last resort
    try:
        for _round in range(MAX_TOOL_ROUNDS):
            r = groq_client.chat.completions.create(
                model=CHAT_MODEL, messages=msgs, max_tokens=600,
                tools=TOOL_SPECS, tool_choice="auto",
                timeout=GROQ_TIMEOUT_S)
            msg = r.choices[0].message

            if not msg.tool_calls:
                reply = clean_latex(strip_leaked_reasoning((msg.content or "").strip()))
                conversation_history[-1] = {"role": "user", "content": user_visible}
                stored = reply
                if actions_log:
                    stored += ("\n[Actions performed this turn: "
                               + "; ".join(actions_log) + "]")
                conversation_history.append({"role": "assistant", "content": stored})
                if len(conversation_history) > 40:
                    del conversation_history[:2]
                return reply

            msgs.append({
                "role": "assistant",
                "content": msg.content or "",
                "tool_calls": [{
                    "id": tc.id, "type": "function",
                    "function": {"name": tc.function.name,
                                 "arguments": tc.function.arguments},
                } for tc in msg.tool_calls],
            })
            for tc in msg.tool_calls:
                if tc.function.name == "view_screen":
                    # UI interaction required — hand control back to the
                    # GUI thread; the snip flow continues the conversation.
                    conversation_history[-1] = {"role": "user", "content": user_visible}
                    return _SNIP_SENTINEL
                try:
                    args = json.loads(tc.function.arguments or "{}")
                except json.JSONDecodeError:
                    args = {}
                key = (tc.function.name,
                       json.dumps(args, sort_keys=True))
                if tc.function.name == "search_web" and search_web_used:
                    # search_web is a last resort — cap it at once per turn
                    # so the model can't escape to a generic search instead
                    # of using web_answers (finds URLs) + open_url (opens
                    # them), even with a different query each call.
                    print(f"[Kogane] search_web capped — already used this turn", flush=True)
                    result = ("search_web already used this turn. Use web_answers "
                              "to find the URL, then open_url.")
                elif key in executed:
                    # Same tool, same arguments, same turn: never re-run
                    # the side effect. Feed the cached result back with an
                    # explicit instruction to answer in text.
                    print(f"[Kogane] duplicate tool call suppressed: {key[0]}", flush=True)
                    result = (f"ALREADY DONE this turn — result was: "
                              f"{executed[key]} Do not call this tool "
                              f"again. Respond to the player in text now.")
                else:
                    result = _run_tool(tc.function.name, args, status_cb, reminder_cb, user_visible)
                    executed[key] = result
                    actions_log.append(f"{tc.function.name}: {result[:100]}")
                    if tc.function.name == "search_web":
                        search_web_used = True
                msgs.append({"role": "tool", "tool_call_id": tc.id,
                             "content": result})

        conversation_history[-1] = {"role": "user", "content": user_visible}
        return "Objective could not be completed within limits."
    except Exception as exc:
        # Never leave the oversized model_content in history on failure
        conversation_history[-1] = {"role": "user", "content": user_visible}
        if _is_rate_limit(exc):
            raise _RateLimitError() from exc
        raise

def _groq_vision(b64: str, prompt: str) -> str:
    msgs = [
        {"role": "system", "content": _system_prompt()},
        {"role": "user", "content": [
            {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}"}},
            {"type": "text", "text": prompt}
        ]},
    ]
    _debug_print_messages("vision", msgs)
    r = groq_client.chat.completions.create(
        model=VISION_MODEL,
        messages=msgs,
        max_tokens=400,
        timeout=GROQ_TIMEOUT_S,
    )
    return clean_latex(r.choices[0].message.content.strip())

def ask_groq_vision(b64: str, prompt: str = VISION_PROMPT) -> str:
    try:
        return _groq_vision(b64, prompt)
    except Exception as exc:
        if _is_rate_limit(exc):
            raise _RateLimitError() from exc
        raise

def ask_groq_with_image(user_input: str, b64: str) -> str:
    try:
        reply = _groq_vision(b64, user_input)
    except Exception as exc:
        if _is_rate_limit(exc):
            raise _RateLimitError() from exc
        raise
    conversation_history.append({"role": "user", "content": user_input})
    conversation_history.append({"role": "assistant", "content": reply})
    return reply

# ── SCREEN CAPTURE ────────────────────────────────────────────────────
def screen_to_b64(region=None) -> str:
    buf = io.BytesIO()
    ImageGrab.grab(bbox=region).save(buf, format="PNG")
    return base64.b64encode(buf.getvalue()).decode()

# ── FILE ATTACHMENTS ──────────────────────────────────────────────────
_TEXT_EXTS = {".txt", ".md", ".py", ".js", ".ts", ".html", ".css", ".json",
              ".csv", ".log", ".xml", ".yaml", ".yml", ".ini", ".cfg",
              ".java", ".c", ".cpp", ".h", ".sh", ".ps1", ".sql", ".lua"}
_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".bmp"}
_MAX_ATTACH_CHARS = 12000

def read_attachment(path: str) -> dict:
    """Returns {"kind": "text"|"image"|"error", "name", "content"|"b64"}."""
    name = os.path.basename(path)
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext in _IMAGE_EXTS:
            img = PILImage.open(path)
            img.thumbnail((1280, 1280))
            buf = io.BytesIO()
            img.convert("RGB").save(buf, format="JPEG", quality=85)
            return {"kind": "image", "name": name,
                    "b64": base64.b64encode(buf.getvalue()).decode()}
        if ext == ".pdf":
            try:
                from pypdf import PdfReader
            except ImportError:
                return {"kind": "error", "name": name,
                        "content": "PDF support requires pypdf — run: pip install pypdf"}
            reader = PdfReader(path)
            pages = [pg.extract_text() or "" for pg in reader.pages[:20]]
            text = "\n".join(pages).strip()
            if not text:
                return {"kind": "error", "name": name,
                        "content": "No extractable text in this PDF (it may be scanned)."}
            if len(text) > _MAX_ATTACH_CHARS:
                text = text[:_MAX_ATTACH_CHARS] + "\n[…truncated…]"
            return {"kind": "text", "name": name, "content": text}
        if ext in _TEXT_EXTS or ext == "":
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read(_MAX_ATTACH_CHARS + 1)
            if len(text) > _MAX_ATTACH_CHARS:
                text = text[:_MAX_ATTACH_CHARS] + "\n[…truncated…]"
            return {"kind": "text", "name": name, "content": text}
        return {"kind": "error", "name": name,
                "content": f"Unsupported file type: {ext}"}
    except Exception as exc:
        return {"kind": "error", "name": name,
                "content": f"Could not read {name}: {_fmt_err(exc)}"}

# ── GROQ WORKER (QObject, moved to a QThread) ────────────────────────
# Lives on a background QThread. run() must NEVER touch a widget — it
# calls the pure network/tool function and emits signals; Qt marshals
# delivery onto the GUI thread. status/reminder follow the same rule.
class GroqWorker(QObject):
    response_ready = pyqtSignal(str)
    error          = pyqtSignal(str)
    status         = pyqtSignal(str)          # "— opened Spotify —" lines
    reminder       = pyqtSignal(int, str)     # ms, message

    def __init__(self, fn):
        super().__init__()
        self._fn = fn

    def run(self):
        try:
            result = self._fn(self.status.emit, self.reminder.emit)
        except Exception as exc:
            msg = _fmt_err(exc) if str(exc) else type(exc).__name__
            print(f"[Kogane] worker error: {msg}", flush=True)
            traceback.print_exc()
            self.error.emit(_user_facing_error(exc))
            return
        self.response_ready.emit(result)

# ── REGION SELECTOR (screen snip overlay) ────────────────────────────
class RegionSelector(QWidget):
    region_selected = pyqtSignal(int, int, int, int)
    cancelled       = pyqtSignal()

    def __init__(self):
        super().__init__()
        self.setWindowFlags(Qt.FramelessWindowHint | Qt.WindowStaysOnTopHint | Qt.Tool)
        self.setAttribute(Qt.WA_TranslucentBackground)
        self.setCursor(Qt.CrossCursor)
        self.setMouseTracking(True)
        self._start = self._end = None
        self._active = False
        self._dim_alpha = 0
        self._ant_off = 0.0
        self._ant_timer = QTimer(self)
        self._ant_timer.timeout.connect(self._tick)

    def getDimAlpha(self): return self._dim_alpha
    def setDimAlpha(self, v): self._dim_alpha = v; self.update()
    dimAlpha = pyqtProperty(int, getDimAlpha, setDimAlpha)

    def activate(self):
        vg = QApplication.primaryScreen().virtualGeometry()
        self.setGeometry(vg)
        self._start = self._end = None
        self._active = False
        self._dim_alpha = 0
        self.show(); self.raise_(); self.activateWindow(); self.setFocus()
        a = QPropertyAnimation(self, b"dimAlpha")
        a.setStartValue(0); a.setEndValue(110)
        a.setDuration(150); a.setEasingCurve(QEasingCurve.OutCubic)
        a.start(); self._anim = a

    def _tick(self):
        self._ant_off = (self._ant_off + 0.8) % 8.0
        if self._start and self._end:
            x1 = min(self._start.x(), self._end.x())
            y1 = min(self._start.y(), self._end.y())
            w  = abs(self._end.x()-self._start.x())
            h  = abs(self._end.y()-self._start.y())
            self.update(QRect(x1-2, y1-2, w+4, h+4))
        else:
            self.update()

    def keyPressEvent(self, e):
        if e.key() == Qt.Key_Escape:
            self._ant_timer.stop(); self.hide(); self.cancelled.emit()

    def mousePressEvent(self, e):
        if e.button() == Qt.LeftButton:
            self._start = e.pos(); self._end = e.pos()
            self._active = True; self._ant_timer.start(25)

    def mouseMoveEvent(self, e):
        if self._active:
            self._end = e.pos(); self.update()

    def mouseReleaseEvent(self, e):
        if e.button() == Qt.LeftButton and self._active:
            self._end = e.pos(); self._active = False
            self._ant_timer.stop(); self.hide()
            if self._start and self._end:
                x1 = min(self._start.x(), self._end.x())
                y1 = min(self._start.y(), self._end.y())
                x2 = max(self._start.x(), self._end.x())
                y2 = max(self._start.y(), self._end.y())
                if x2-x1 > 4 and y2-y1 > 4:
                    self.region_selected.emit(x1, y1, x2, y2)
                    return
            self.cancelled.emit()

    def paintEvent(self, _):
        p = QPainter(self)
        p.setRenderHint(QPainter.Antialiasing)
        p.fillRect(self.rect(), QColor(0, 0, 0, self._dim_alpha))
        if self._start and self._end:
            x1 = min(self._start.x(), self._end.x())
            y1 = min(self._start.y(), self._end.y())
            w  = abs(self._end.x()-self._start.x())
            h  = abs(self._end.y()-self._start.y())
            sel = QRect(x1, y1, w, h)
            p.setCompositionMode(QPainter.CompositionMode_Clear)
            p.fillRect(sel, Qt.transparent)
            p.setCompositionMode(QPainter.CompositionMode_SourceOver)
            for col, off in ((QColor(255,255,255,160), self._ant_off),
                             (QColor(200,144,28,210),  self._ant_off+4.0)):
                pen = QPen(col, 1.5, Qt.CustomDashLine)
                pen.setDashPattern([4.0,4.0]); pen.setDashOffset(off)
                p.setPen(pen); p.setBrush(Qt.NoBrush); p.drawRect(sel)
            p.setPen(QColor(255,255,255,180))
            p.setFont(_px_font(FS_CAPTION))
            ly = y1-8 if y1>20 else y1+h+16
            p.drawText(x1+4, ly, f"{w} × {h}")
        else:
            p.setPen(QPen(QColor(255,255,255,80), 1))
            p.setFont(_px_font(FS_CAPTION))
            p.drawText(self.rect(), Qt.AlignCenter,
                       "Click and drag to select  ·  Esc to cancel")
        p.end()

# ── AVATAR ────────────────────────────────────────────────────────────
_avatar_cache: dict = {}

def _circular_pixmap(path: str, size: int) -> QPixmap:
    key = (path, size)
    cached = _avatar_cache.get(key)
    if cached is not None:
        return cached
    src = QPixmap(path)
    result = QPixmap(size, size)
    result.fill(Qt.transparent)
    if not src.isNull():
        src = src.scaled(size, size, Qt.KeepAspectRatioByExpanding, Qt.SmoothTransformation)
        painter = QPainter(result)
        painter.setRenderHint(QPainter.Antialiasing)
        clip = QPainterPath()
        clip.addEllipse(0.0, 0.0, float(size), float(size))
        painter.setClipPath(clip)
        painter.drawPixmap((size - src.width()) // 2, (size - src.height()) // 2, src)
        painter.end()
    _avatar_cache[key] = result
    return result

# ── TEXT HELPERS ──────────────────────────────────────────────────────
_ZWS = "​"
_LONG_TOKEN_RE = re.compile(r'\S{25,}')

def _soften_long_tokens(text: str) -> str:
    return _LONG_TOKEN_RE.sub(
        lambda m: _ZWS.join(m.group(0)[i:i+24] for i in range(0, len(m.group(0)), 24)),
        text
    )

def _rich_text(text: str) -> str:
    """Escape + convert to rich text with a proper line height (QSS has no
    line-height for QLabel; rich text blocks do)."""
    safe = html.escape(_soften_long_tokens(text)).replace("\n", "<br>")
    return f'<p style="margin:0; line-height:145%;">{safe}</p>'

# ── BUTTON SYSTEM ────────────────────────────────────────────────────
# Three classes. Radius is always derived from the height it pairs with,
# never hardcoded independently (radius > height/2 breaks Qt corners).
ICON_BTN = 40          # header circles
SEND_BTN = 42          # circular send

class IconButton(QPushButton):
    """Icon button — paints a crisp vector glyph via QPainterPath instead
    of a font/emoji character (font glyphs sit off-center and look
    thin/blurry at small sizes; same reasoning as SendButton)."""

    _KINDS = ("scissors", "plus", "clock", "chevron_down", "compose")

    def __init__(self, kind: str, tooltip: str, size: int = ICON_BTN, parent=None):
        assert kind in self._KINDS, f"unknown IconButton kind: {kind}"
        super().__init__("", parent)
        self._kind = kind
        self._size = size
        self.setCursor(Qt.PointingHandCursor)
        self.setFixedSize(size, size)
        self.setToolTip(tooltip)
        self.setAccessibleName(tooltip)
        self.setStyleSheet(
            f"QPushButton {{"
            f"  background: transparent; border: none;"
            f"  border-radius: {size // 2}px;"
            f"}}"
            f"QPushButton:hover {{ background: {P.surface}; }}"
            f"QPushButton:pressed {{ background: {P.surface_hover}; }}"
            f"QPushButton:checked {{ background: {P.surface}; }}"
        )

    def _build_path(self, cx: float, cy: float) -> QPainterPath:
        path = QPainterPath()
        if self._kind == "scissors":
            path.addEllipse(QPointF(cx - 4.5, cy + 4.5), 2.3, 2.3)
            path.addEllipse(QPointF(cx + 4.5, cy + 4.5), 2.3, 2.3)
            path.moveTo(cx - 4.5, cy + 4.5)
            path.lineTo(cx + 5.5, cy - 5.5)
            path.moveTo(cx + 4.5, cy + 4.5)
            path.lineTo(cx - 5.5, cy - 5.5)
        elif self._kind == "plus":
            path.moveTo(cx, cy - 7)
            path.lineTo(cx, cy + 7)
            path.moveTo(cx - 7, cy)
            path.lineTo(cx + 7, cy)
        elif self._kind == "clock":
            path.addEllipse(QPointF(cx, cy), 7.5, 7.5)
            path.moveTo(cx, cy)
            path.lineTo(cx, cy - 4.5)          # hour hand, up
            path.moveTo(cx, cy)
            path.lineTo(cx + 4.0, cy - 1.0)     # minute hand, toward 2 o'clock
        elif self._kind == "chevron_down":
            path.moveTo(cx - 5.5, cy - 2.5)
            path.lineTo(cx, cy + 3.5)
            path.lineTo(cx + 5.5, cy - 2.5)
        elif self._kind == "compose":
            path.addRect(cx - 6.5, cy - 1.0, 7.0, 7.0)   # page
            path.moveTo(cx - 3.0, cy + 2.5)
            path.lineTo(cx + 6.5, cy - 7.0)               # pencil stroke
        return path

    def paintEvent(self, event):
        super().paintEvent(event)          # QSS circle background / hover fill
        p = QPainter(self)
        p.setRenderHint(QPainter.Antialiasing)
        cx, cy = self.width() / 2.0, self.height() / 2.0
        active = self.underMouse() or self.isChecked()
        color = QColor(P.text if active else P.text_muted)
        pen = QPen(color, 2.0)
        pen.setCapStyle(Qt.RoundCap)
        pen.setJoinStyle(Qt.RoundJoin)
        p.setPen(pen)
        p.setBrush(Qt.NoBrush)
        p.drawPath(self._build_path(cx, cy))
        p.end()

def pill_metrics(font_px: int, pad_v: int = 12) -> tuple:
    h = _line_h(font_px) + pad_v * 2
    return h, h // 2

# ── EMPTY STATE ──────────────────────────────────────────────────────
class EmptyState(QWidget):
    suggestion_clicked = pyqtSignal(str)

    _PROMPTS = [
        "what's on my screen?",
        "play asake on spotify",
        "remind me in 20 minutes to stretch",
    ]

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setAttribute(Qt.WA_TranslucentBackground)
        layout = QVBoxLayout(self)
        layout.setAlignment(Qt.AlignVCenter)
        layout.setContentsMargins(18, 36, 18, 28)
        layout.setSpacing(0)

        purpose = QLabel("ask, see, open.")
        purpose.setAlignment(Qt.AlignCenter)
        purpose.setStyleSheet(
            f"color: {P.text_muted}; font: 600 {FS_SMALL}px {_font_stack()};"
        )
        layout.addWidget(purpose)
        layout.addSpacing(22)

        chip_h, chip_r = pill_metrics(FS_CHIP)
        for text in self._PROMPTS:
            btn = QPushButton(text)
            btn.setCursor(Qt.PointingHandCursor)
            btn.setFixedHeight(chip_h)
            btn.setAccessibleName(f"Try: {text}")
            btn.setStyleSheet(
                f"QPushButton {{"
                f"  background: {P.surface};"
                f"  color: {P.text};"
                f"  border: 1px solid {P.surface_border};"
                f"  border-radius: {chip_r}px;"
                f"  font: {FS_CHIP}px {_font_stack_medium()};"
                f"  text-align: left; padding: 0 20px;"
                f"}}"
                f"QPushButton:hover {{"
                f"  background: {P.surface_hover}; border-color: {P.text_muted};"
                f"}}"
                f"QPushButton:pressed {{"
                f"  background: {P.user_bubble}; border-color: {_rgba(P.accent, 110)};"
                f"}}"
            )
            btn.clicked.connect(
                lambda checked=False, t=text: self.suggestion_clicked.emit(t)
            )
            layout.addWidget(btn)
            layout.addSpacing(8)

        layout.addSpacing(22)
        shortcut = QLabel("Ctrl+Shift+S to snip your screen")
        shortcut.setAlignment(Qt.AlignCenter)
        shortcut.setStyleSheet(
            f"color: {P.muted_dim}; font: {FS_SMALL}px {_font_stack()};"
        )
        layout.addWidget(shortcut)

# ── ENTRANCE WRAPPER ─────────────────────────────────────────────────
class _SlideInRow(QWidget):
    def __init__(self, inner: QWidget, parent=None):
        super().__init__(parent)
        self.setAttribute(Qt.WA_TranslucentBackground)
        self._offset = 12
        self._layout = QVBoxLayout(self)
        self._layout.setContentsMargins(0, self._offset, 0, 0)
        self._layout.setSpacing(0)
        self._layout.addWidget(inner)

    def _get_offset(self): return self._offset
    def _set_offset(self, v):
        self._offset = v
        self._layout.setContentsMargins(0, v, 0, 0)
    offset = pyqtProperty(int, _get_offset, _set_offset)

# ── MESSAGE BUBBLE ────────────────────────────────────────────────────
_BUBBLE_MAX_W = int(PANEL_W * 0.78)

class MessageBubble(QWidget):
    def __init__(self, role: str, text: str, parent=None):
        super().__init__(parent)
        self.setAttribute(Qt.WA_TranslucentBackground)
        row = QHBoxLayout(self)
        row.setContentsMargins(0, 0, 0, 0)
        row.setSpacing(8)

        label = QLabel()
        label.setTextFormat(Qt.RichText)
        label.setText(_rich_text(text))
        label.setWordWrap(True)
        label.setTextInteractionFlags(Qt.TextSelectableByMouse)
        label.setMaximumWidth(_BUBBLE_MAX_W)
        label.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Minimum)

        if role == "user":
            label.setStyleSheet(
                f"background: {P.user_bubble}; color: {P.user_text};"
                f"border: 1px solid {P.user_border};"
                f"border-top-left-radius: 16px; border-top-right-radius: 16px;"
                f"border-bottom-left-radius: 16px; border-bottom-right-radius: 4px;"
                f"padding: 12px 16px; font-family: {_font_stack_medium()};"
                f"font-size: {FS_BODY}px;"
            )
            row.addStretch()
            row.addWidget(label)
        else:
            avatar = QLabel()
            avatar.setPixmap(_circular_pixmap(ICON_PATH, 24))
            avatar.setFixedSize(24, 24)
            row.addWidget(avatar, 0, Qt.AlignBottom)
            label.setStyleSheet(
                f"background: {P.surface}; color: {P.text};"
                f"border-top-left-radius: 16px; border-top-right-radius: 16px;"
                f"border-bottom-right-radius: 16px; border-bottom-left-radius: 4px;"
                f"padding: 12px 16px; font-family: {_font_stack_medium()};"
                f"font-size: {FS_BODY}px;"
            )
            row.addWidget(label)
            row.addStretch()

class StatusLine(QWidget):
    """Muted, centered, italic — how tool actions appear in the chat."""
    def __init__(self, text: str, parent=None):
        super().__init__(parent)
        self.setAttribute(Qt.WA_TranslucentBackground)
        lay = QHBoxLayout(self)
        lay.setContentsMargins(0, 2, 0, 2)
        lbl = QLabel(f"—\u00A0{text}\u00A0—")
        lbl.setAlignment(Qt.AlignCenter)
        lbl.setWordWrap(True)
        lbl.setStyleSheet(
            f"color: {P.text_muted}; font: italic {FS_CAPTION}px {_font_stack()};"
            f"background: transparent;"
        )
        lay.addStretch(); lay.addWidget(lbl); lay.addStretch()

# ── TYPING INDICATOR ─────────────────────────────────────────────────
class StarTypingIndicator(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setAttribute(Qt.WA_TranslucentBackground)
        self.setFixedSize(46, 16)
        self._phase = 0.0
        self._clock = QElapsedTimer()
        self._timer = QTimer(self)
        self._timer.setInterval(16)
        self._timer.timeout.connect(self._tick)

    def start(self):
        self._clock.start(); self._timer.start(); self.show()

    def stop(self):
        self._timer.stop(); self.hide()

    def _tick(self):
        self._phase = self._clock.elapsed() / 1000.0
        self.update()

    def paintEvent(self, _):
        p = QPainter(self)
        p.setRenderHint(QPainter.Antialiasing)
        cy = self.height() / 2.0
        for i in range(3):
            cx = 9 + i * 15
            t = (self._phase * 1.5 - i * 0.3) % 1.0
            op = 0.25 + 0.75 * 0.5 * (1.0 - math.cos(2.0 * math.pi * t))
            color = QColor(P.accent)
            color.setAlphaF(op)
            p.setBrush(color); p.setPen(Qt.NoPen)
            p.drawPath(_StarParticle.star_path(cx, cy, 4.5, 0.0))
        p.end()

class TypingBubble(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setAttribute(Qt.WA_TranslucentBackground)
        row = QHBoxLayout(self)
        row.setContentsMargins(0, 0, 0, 0)
        row.setSpacing(8)
        avatar = QLabel()
        avatar.setPixmap(_circular_pixmap(ICON_PATH, 24))
        avatar.setFixedSize(24, 24)
        row.addWidget(avatar, 0, Qt.AlignBottom)
        bubble = QFrame()
        bubble.setStyleSheet(
            f"background: {P.surface};"
            f"border-top-left-radius: 16px; border-top-right-radius: 16px;"
            f"border-bottom-right-radius: 16px; border-bottom-left-radius: 4px;"
        )
        b_l = QHBoxLayout(bubble)
        b_l.setContentsMargins(16, 14, 16, 14)
        self._stars = StarTypingIndicator()
        b_l.addWidget(self._stars)
        row.addWidget(bubble)
        row.addStretch()

    def start(self): self._stars.start()
    def stop(self):  self._stars.stop()

# ── SEND BUTTON (painted arrow — no font glyph) ──────────────────────
class SendButton(QPushButton):
    """42px gold circle. Paints a crisp arrow with QPainterPath instead of
    the "↑" glyph (font glyphs sit off-center and look thin). While Kogane
    is thinking it paints a pulsing star instead — a busy indicator, not a
    stop control."""

    def __init__(self, parent=None):
        super().__init__("", parent)
        self.setCursor(Qt.PointingHandCursor)
        self.setFixedSize(SEND_BTN, SEND_BTN)
        self.setAccessibleName("Send message")
        self.setStyleSheet(
            f"QPushButton {{"
            f"  background: {P.accent};"
            f"  border-radius: {SEND_BTN // 2}px; border: none;"
            f"}}"
            f"QPushButton:hover {{ background: {P.accent_hover}; }}"
            f"QPushButton:pressed {{ background: {P.accent_press}; }}"
            f"QPushButton:disabled {{ background: {P.surface}; }}"
        )
        self._thinking = False
        self._clock = QElapsedTimer()
        self._pulse = QTimer(self)
        self._pulse.setInterval(16)
        self._pulse.timeout.connect(self.update)

    def set_thinking(self, on: bool):
        self._thinking = on
        if on:
            self._clock.start()
            self._pulse.start()
        else:
            self._pulse.stop()
        self.update()

    def paintEvent(self, event):
        super().paintEvent(event)          # QSS circle background
        p = QPainter(self)
        p.setRenderHint(QPainter.Antialiasing)
        cx, cy = self.width() / 2.0, self.height() / 2.0
        if self._thinking:
            t = (self._clock.elapsed() / 1000.0 * 1.4) % 1.0
            op = 0.35 + 0.65 * 0.5 * (1.0 - math.cos(2.0 * math.pi * t))
            color = QColor(P.muted_dim if not self.isEnabled() else P.accent_text)
            color = QColor(P.text_muted)
            color.setAlphaF(op)
            p.setPen(Qt.NoPen)
            p.setBrush(color)
            p.drawPath(_StarParticle.star_path(cx, cy, 7.0,
                                               self._clock.elapsed() / 25.0))
        else:
            color = QColor(P.accent_text if self.isEnabled() else P.muted_dim)
            pen = QPen(color, 2.5)
            pen.setCapStyle(Qt.RoundCap)
            pen.setJoinStyle(Qt.RoundJoin)
            p.setPen(pen)
            path = QPainterPath()
            path.moveTo(cx, cy + 6)          # shaft bottom
            path.lineTo(cx, cy - 6)          # shaft top
            path.moveTo(cx - 5.5, cy - 0.5)  # chevron left
            path.lineTo(cx, cy - 6)
            path.lineTo(cx + 5.5, cy - 0.5)  # chevron right
            p.drawPath(path)
        p.end()


# ── PANEL FRAME (Escape fallback) ────────────────────────────────────
class _PanelFrame(QFrame):
    escape_pressed = pyqtSignal()

    def keyPressEvent(self, e):
        if e.key() == Qt.Key_Escape:
            self.escape_pressed.emit()
            return
        super().keyPressEvent(e)

# ── MAIN WINDOW ───────────────────────────────────────────────────────
class KoganeWindow(QWidget):
    _trigger_snip_sig = pyqtSignal()
    _toggle_sig       = pyqtSignal()
    _new_conv_sig     = pyqtSignal()
    _history_sig      = pyqtSignal()
    _quit_sig         = pyqtSignal()

    _DEFAULT_BOB_SPEED  = kogane_animation.BOB_SPEED
    _THINKING_BOB_SPEED = kogane_animation.BOB_SPEED * 1.8

    _MAX_INPUT = 4000

    def __init__(self):
        super().__init__()
        self._chat_open       = False
        self._locked          = False
        self._last_vision_b64 = None
        self._attachment: dict | None = None

        self._thread: QThread | None      = None
        self._worker: GroqWorker | None   = None
        self._title_jobs: list            = []   # keeps background title threads alive
        self._panel_anim: QParallelAnimationGroup | None = None
        self._send_btn_base_geo: QRect | None = None
        self._send_scale_anim: QPropertyAnimation | None = None

        self.setWindowFlags(
            Qt.FramelessWindowHint | Qt.WindowStaysOnTopHint | Qt.Tool
        )
        self.setAttribute(Qt.WA_TranslucentBackground)
        self.setAttribute(Qt.WA_NoSystemBackground)

        self._build_ui()
        self._connect_signals()

        self.setFixedSize(PANEL_W, PANEL_H)
        self.face_widget.place_bottom_right()
        self.face_widget.summon()
        self._reschedule_saved_reminders()

    # ── UI CONSTRUCTION ───────────────────────────────────────────────
    def _build_ui(self):
        self._root = QVBoxLayout(self)
        self._root.setContentsMargins(0, 0, 0, 0)
        self._root.setSpacing(0)

        self._panel = self._make_panel()
        # NOTE: deliberately NO QGraphicsOpacityEffect on the panel. A
        # panel-level effect nested with per-row bubble effects is
        # unsupported by Qt and caused blank bubbles + QPainter spam.
        # The panel fade uses setWindowOpacity (top-level window) instead.
        self._panel.hide()
        self._root.addWidget(self._panel)

        self.face_widget = KoganeIconWidget(ICON_PATH)
        self.face_widget.clicked.connect(self._toggle_chat)

        self._selector = RegionSelector()
        self._selector.region_selected.connect(self._on_region_selected)
        self._selector.cancelled.connect(self._on_snip_cancelled)

    def _make_panel(self) -> QFrame:
        panel = _PanelFrame()
        panel.setObjectName("koganePanel")
        panel.setStyleSheet(
            f"#koganePanel {{"
            f"  background: {P.bg};"
            f"  border-radius: 18px;"
            f"  border: 1px solid {P.divider};"
            f"}}"
        )
        panel.escape_pressed.connect(self._on_escape)
        esc = QShortcut(QKeySequence(Qt.Key_Escape), panel)
        esc.setContext(Qt.WidgetWithChildrenShortcut)
        esc.activated.connect(self._on_escape)

        layout = QVBoxLayout(panel)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        layout.addWidget(self._make_header())

        div = QFrame(); div.setFixedHeight(1)
        div.setStyleSheet(f"background: {P.divider}; border: none;")
        layout.addWidget(div)

        self._stack = QStackedWidget()
        self._stack.addWidget(self._make_chat_page())     # page 0
        self._stack.addWidget(self._make_history_page())  # page 1
        layout.addWidget(self._stack, 1)
        return panel

    def _make_header(self) -> QWidget:
        hdr = QWidget()
        hdr.setFixedHeight(60)
        hdr.setStyleSheet("background: transparent;")
        h = QHBoxLayout(hdr)
        h.setContentsMargins(14, 0, 10, 0)
        h.setSpacing(10)

        avatar = QLabel()
        avatar.setPixmap(_circular_pixmap(ICON_PATH, 32))
        avatar.setFixedSize(32, 32)
        h.addWidget(avatar)

        titles = QVBoxLayout()
        titles.setSpacing(1)
        titles.setContentsMargins(0, 0, 0, 0)
        title = QLabel("Kogane")
        title.setStyleSheet(
            f"color: {P.text}; font: {FS_TITLE}px {_font_stack_semibold()};"
        )
        status = QLabel("Online")
        status.setStyleSheet(
            f"color: {P.text_muted}; font: {FS_SMALL}px {_font_stack()};"
        )
        titles.addStretch(1)
        titles.addWidget(title)
        titles.addWidget(status)
        titles.addStretch(1)
        h.addLayout(titles)
        h.addStretch()

        snip_btn = IconButton("scissors", "Screen snip (Ctrl+Shift+S)")
        snip_btn.clicked.connect(self._start_snip)
        h.addWidget(snip_btn)

        new_btn = IconButton("compose", "New conversation")
        new_btn.clicked.connect(self._new_conversation)
        h.addWidget(new_btn)

        self._hist_btn = IconButton("clock", "History")
        self._hist_btn.setCheckable(True)
        self._hist_btn.clicked.connect(self._toggle_history)
        h.addWidget(self._hist_btn)

        min_btn = IconButton("chevron_down", "Minimize (Esc)")
        min_btn.clicked.connect(self._close_panel_animated)
        h.addWidget(min_btn)
        return hdr

    # ── CHAT PAGE ─────────────────────────────────────────────────────
    def _make_chat_page(self) -> QWidget:
        page = QWidget()
        page.setStyleSheet("background: transparent;")
        layout = QVBoxLayout(page)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        self._scroll = QScrollArea()
        self._scroll.setWidgetResizable(True)
        self._scroll.setFrameShape(QFrame.NoFrame)
        self._scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self._scroll.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self._scroll.setStyleSheet(
            f"QScrollArea {{ background: transparent; border: none; }}"
            f"QScrollBar:vertical {{"
            f"  background: transparent; width: 6px; margin: 4px 2px;"
            f"}}"
            f"QScrollBar::handle:vertical {{"
            f"  background: {P.surface_hover}; border-radius: 3px; min-height: 24px;"
            f"}}"
            f"QScrollBar::add-line, QScrollBar::sub-line {{ height: 0; }}"
        )

        msg_container = QWidget()
        msg_container.setStyleSheet("background: transparent;")
        self._msg_container = msg_container
        self._msg_layout = QVBoxLayout(msg_container)
        self._msg_layout.setContentsMargins(16, 12, 16, 12)
        self._msg_layout.setSpacing(12)

        self._empty_state = EmptyState()
        self._empty_state.suggestion_clicked.connect(self._fill_input)
        self._msg_layout.addWidget(self._empty_state)
        self._msg_layout.addStretch(1)

        self._thinking_wrap = QWidget()
        self._thinking_wrap.setStyleSheet("background: transparent;")
        tw = QVBoxLayout(self._thinking_wrap)
        tw.setContentsMargins(0, 0, 0, 0)
        self._typing_bubble = TypingBubble()
        tw.addWidget(self._typing_bubble)
        self._thinking_wrap.hide()
        self._msg_layout.addWidget(self._thinking_wrap)

        self._scroll.setWidget(msg_container)
        layout.addWidget(self._scroll, 1)

        # attachment chip row (hidden until a file is attached)
        self._attach_row = QWidget()
        self._attach_row.setStyleSheet("background: transparent;")
        ar = QHBoxLayout(self._attach_row)
        ar.setContentsMargins(16, 0, 16, 4)
        self._attach_chip = QPushButton()
        self._attach_chip.setCursor(Qt.PointingHandCursor)
        chip_h, chip_r = pill_metrics(FS_SMALL, pad_v=6)
        self._attach_chip.setFixedHeight(chip_h)
        self._attach_chip.setStyleSheet(
            f"QPushButton {{"
            f"  background: {P.surface}; color: {P.text_muted};"
            f"  border: 1px solid {P.surface_border};"
            f"  border-radius: {chip_r}px; padding: 0 14px;"
            f"  font: {FS_SMALL}px {_font_stack()};"
            f"}}"
            f"QPushButton:hover {{ color: {P.danger}; border-color: {P.danger}; }}"
        )
        self._attach_chip.clicked.connect(self._clear_attachment)
        ar.addWidget(self._attach_chip)
        ar.addStretch()
        self._attach_row.hide()
        layout.addWidget(self._attach_row)

        layout.addWidget(self._make_input_bar())
        return page

    def _make_input_bar(self) -> QWidget:
        bar = QWidget()
        bar.setStyleSheet("background: transparent;")
        l = QHBoxLayout(bar)
        l.setContentsMargins(16, 8, 16, 16)
        l.setSpacing(8)

        # Input height derives from font metrics — text can never clip.
        pad_v = 11
        input_h = _line_h(FS_BODY) + pad_v * 2 + 2
        input_r = input_h // 2

        attach_btn = IconButton("plus", "Attach a file", size=input_h)
        attach_btn.clicked.connect(self._pick_attachment)
        l.addWidget(attach_btn, 0, Qt.AlignVCenter)

        self._input = QTextEdit()
        self._input.setPlaceholderText("Ask anything…")
        self._input.setAcceptRichText(False)
        self._input.setFixedHeight(input_h)
        self._input.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self._input.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self._input.document().setDocumentMargin(0)
        self._input.setFont(_px_font(FS_BODY, family=FONT_FAMILY_MEDIUM or FONT_FAMILY))
        self._input.setStyleSheet(
            f"QTextEdit {{"
            f"  background: {P.surface}; color: {P.text};"
            f"  border: 1px solid {P.surface_border};"
            f"  border-radius: {input_r}px;"
            f"  padding: {pad_v - 1}px 18px;"
            f"  font-size: {FS_BODY}px; font-family: {_font_stack_medium()};"
            f"  selection-background-color: {_rgba(P.accent, 90)};"
            f"}}"
            f"QTextEdit:focus {{ border: 1px solid {P.accent}; }}"
        )
        # Gold caret so the typing position is findable
        pal = self._input.palette()
        pal.setColor(pal.Text, QColor(P.text))
        self._input.setPalette(pal)
        self._input.installEventFilter(self)
        l.addWidget(self._input, 1)

        send_btn = SendButton()
        send_btn.clicked.connect(self._on_send)
        send_btn.pressed.connect(self._on_send_btn_pressed)
        send_btn.released.connect(self._on_send_btn_released)
        self._send_btn = send_btn
        l.addWidget(send_btn, 0, Qt.AlignVCenter)
        return bar

    # ── HISTORY PAGE ──────────────────────────────────────────────────
    def _make_history_page(self) -> QWidget:
        page = QWidget()
        page.setStyleSheet("background: transparent;")
        layout = QVBoxLayout(page)
        layout.setContentsMargins(0, 0, 0, 0)

        self._hist_scroll = QScrollArea()
        self._hist_scroll.setWidgetResizable(True)
        self._hist_scroll.setFrameShape(QFrame.NoFrame)
        self._hist_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self._hist_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self._hist_scroll.setStyleSheet(self._scroll.styleSheet())

        self._hist_container = QWidget()
        self._hist_container.setStyleSheet("background: transparent;")
        self._hist_layout = QVBoxLayout(self._hist_container)
        self._hist_layout.setContentsMargins(16, 12, 16, 12)
        self._hist_layout.setSpacing(8)
        self._hist_layout.addStretch(1)

        self._hist_scroll.setWidget(self._hist_container)
        layout.addWidget(self._hist_scroll, 1)
        return page

    @staticmethod
    def _relative_date(stamp: str) -> str:
        try:
            dt = datetime.strptime(stamp.split(".")[0], "%Y-%m-%d %H:%M:%S")
        except Exception:
            return ""
        today = datetime.now().date()
        if dt.date() == today:
            return "today"
        if dt.date() == today - timedelta(days=1):
            return "yesterday"
        return dt.strftime("%b %-d") if os.name != "nt" else dt.strftime("%b %#d")

    def _refresh_history(self):
        while self._hist_layout.count() > 1:
            item = self._hist_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        convs = db_load_conversations()
        if not convs:
            empty = QLabel("No past conversations.")
            empty.setAlignment(Qt.AlignCenter)
            empty.setStyleSheet(
                f"color: {P.text_muted}; font: {FS_CHIP}px {_font_stack()};"
                f"padding: 40px 0;"
            )
            self._hist_layout.insertWidget(0, empty)
            return

        for i, conv in enumerate(convs):
            self._hist_layout.insertWidget(i, self._make_history_row(conv))

    def _make_history_row(self, conv: dict) -> QWidget:
        row = QFrame()
        row.setCursor(Qt.PointingHandCursor)
        row.setStyleSheet(
            f"QFrame {{ background: transparent; border-radius: 12px; }}"
            f"QFrame:hover {{ background: {P.surface}; }}"
        )
        l = QHBoxLayout(row)
        l.setContentsMargins(12, 10, 8, 10)

        texts = QVBoxLayout()
        texts.setSpacing(2)
        title = QLabel(conv["title"])
        title.setStyleSheet(
            f"color: {P.text}; font: {FS_CHIP}px {_font_stack_medium()};"
            f"background: transparent;"
        )
        date = QLabel(self._relative_date(conv.get("updated_at", "")))
        date.setStyleSheet(
            f"color: {P.text_muted}; font: {FS_CAPTION}px {_font_stack()};"
            f"background: transparent;"
        )
        texts.addWidget(title)
        texts.addWidget(date)
        l.addLayout(texts, 1)

        del_btn = QPushButton("✕")
        del_btn.setCursor(Qt.PointingHandCursor)
        del_btn.setFixedSize(28, 28)
        del_btn.setAccessibleName("Delete conversation")
        del_btn.setStyleSheet(
            f"QPushButton {{"
            f"  background: transparent; color: {P.text_muted};"
            f"  border: none; border-radius: 14px; font-size: 13px;"
            f"}}"
            f"QPushButton:hover {{ color: {P.danger}; background: {P.surface_hover}; }}"
        )
        del_btn._armed = False

        def on_delete():
            if not del_btn._armed:
                del_btn._armed = True
                del_btn.setText("sure?")
                del_btn.setFixedSize(52, 28)
                QTimer.singleShot(3000, disarm)
                return
            db_delete_conversation(conv["id"])
            global _active_conv_id
            if _active_conv_id == conv["id"]:
                self._new_conversation()
            self._refresh_history()

        def disarm():
            try:
                if del_btn._armed:
                    del_btn._armed = False
                    del_btn.setText("✕")
                    del_btn.setFixedSize(28, 28)
            except RuntimeError:
                pass

        del_btn.clicked.connect(on_delete)
        l.addWidget(del_btn)

        def open_conv(event):
            if event.button() == Qt.LeftButton:
                self._load_conversation(conv["id"])
        row.mouseReleaseEvent = open_conv
        return row

    def _load_conversation(self, conv_id: int):
        global _active_conv_id
        msgs = db_load_messages(conv_id)
        self._clear_chat_view()
        _active_conv_id = conv_id
        conversation_history.clear()
        for m in msgs:
            role, content = m["role"], m["content"]
            if role == "status":
                self._add_status(content, persist=False, animate=False)
            elif role == "user":
                conversation_history.append({"role": "user", "content": content})
                self._add_message("user", content, animate=False)
            else:
                conversation_history.append({"role": "assistant", "content": content})
                self._add_message("ai", content, animate=False)
        if len(conversation_history) > 40:
            del conversation_history[:len(conversation_history) - 40]
        self._show_chat_page()
        QTimer.singleShot(30, self._scroll_to_bottom)

    def _toggle_history(self):
        if self._stack.currentIndex() == 1:
            self._show_chat_page()
        else:
            self._refresh_history()
            self._stack.setCurrentIndex(1)
            self._hist_btn.setChecked(True)

    def _show_chat_page(self):
        self._stack.setCurrentIndex(0)
        self._hist_btn.setChecked(False)
        self._input.setFocus()

    def _connect_signals(self):
        self._trigger_snip_sig.connect(self._start_snip)
        self._toggle_sig.connect(self._toggle_visibility)
        self._new_conv_sig.connect(self._new_conversation)
        self._history_sig.connect(self._open_history_from_tray)
        self._quit_sig.connect(self._do_quit)

    def _open_history_from_tray(self):
        if not self._chat_open:
            self._open_panel_animated()
        self._refresh_history()
        self._stack.setCurrentIndex(1)
        self._hist_btn.setChecked(True)

    # ── ASYNC HELPER (thread-safe Groq calls) ─────────────────────────
    def _run_async(self, fn, on_result, on_error):
        self._cleanup_worker()

        thread = QThread(self)
        worker = GroqWorker(fn)
        worker.moveToThread(thread)

        thread.started.connect(worker.run)
        worker.response_ready.connect(on_result)
        worker.error.connect(on_error)
        worker.status.connect(self._on_tool_status)
        worker.reminder.connect(self._on_reminder_requested)
        worker.response_ready.connect(thread.quit)
        worker.error.connect(thread.quit)
        worker.response_ready.connect(worker.deleteLater)
        worker.error.connect(worker.deleteLater)
        thread.finished.connect(thread.deleteLater)

        self._thread = thread
        self._worker = worker
        thread.start()

    def _cleanup_worker(self):
        if self._thread is not None:
            try:
                if self._thread.isRunning():
                    self._thread.quit()
                    self._thread.wait(2000)
            except RuntimeError:
                pass
        self._thread = None
        self._worker = None

    # ── PANEL OPEN / CLOSE (windowOpacity — never a graphics effect) ──
    def _panel_target_pos(self) -> QPoint:
        icon_geo = self.face_widget.geometry()
        pad = (WIDGET_SIZE - ICON_SIZE) // 2
        anchor_x = icon_geo.right()  - pad
        anchor_y = icon_geo.top()    + pad
        new_x = anchor_x - PANEL_W
        new_y = anchor_y - PANEL_H - 6
        screen = QApplication.primaryScreen().availableGeometry()
        new_x = max(screen.left(), min(new_x, screen.right()  - PANEL_W))
        new_y = max(screen.top(),  min(new_y, screen.bottom() - PANEL_H))
        return QPoint(new_x, new_y)

    def _open_panel_animated(self):
        if self._chat_open:
            return
        if (self._panel_anim
                and self._panel_anim.state() == QAbstractAnimation.Running):
            return
        self._chat_open = True

        final_pos = self._panel_target_pos()
        start_pos = QPoint(final_pos.x(), final_pos.y() + 24)

        self.setWindowOpacity(0.0)
        self._panel.show()
        self.move(start_pos)
        self.show()
        self.raise_()
        self._input.setFocus()

        if not self.face_widget.isVisible():
            self.face_widget.summon()

        pos_anim = QPropertyAnimation(self, b"pos")
        pos_anim.setStartValue(start_pos)
        pos_anim.setEndValue(final_pos)
        pos_anim.setDuration(220)
        pos_anim.setEasingCurve(QEasingCurve.OutCubic)

        fade_anim = QPropertyAnimation(self, b"windowOpacity")
        fade_anim.setStartValue(0.0)
        fade_anim.setEndValue(1.0)
        fade_anim.setDuration(220)
        fade_anim.setEasingCurve(QEasingCurve.OutCubic)

        group = QParallelAnimationGroup(self)
        group.addAnimation(pos_anim)
        group.addAnimation(fade_anim)
        group.start()
        self._panel_anim = group

    def _close_panel_animated(self):
        if not self._chat_open:
            return
        if (self._panel_anim
                and self._panel_anim.state() == QAbstractAnimation.Running):
            return
        self._chat_open = False

        start_pos = self.pos()
        end_pos = QPoint(start_pos.x(), start_pos.y() + 16)

        pos_anim = QPropertyAnimation(self, b"pos")
        pos_anim.setStartValue(start_pos)
        pos_anim.setEndValue(end_pos)
        pos_anim.setDuration(160)
        pos_anim.setEasingCurve(QEasingCurve.OutCubic)

        fade_anim = QPropertyAnimation(self, b"windowOpacity")
        fade_anim.setStartValue(self.windowOpacity())
        fade_anim.setEndValue(0.0)
        fade_anim.setDuration(160)
        fade_anim.setEasingCurve(QEasingCurve.OutCubic)

        group = QParallelAnimationGroup(self)
        group.addAnimation(pos_anim)
        group.addAnimation(fade_anim)

        def _after():
            self._panel.hide()
            self.hide()
            self.setWindowOpacity(1.0)
            # Icon stays — permanent anchor; only _do_quit dismisses it.

        group.finished.connect(_after)
        group.start()
        self._panel_anim = group

    def _toggle_chat(self):
        if self._chat_open:
            self._close_panel_animated()
        else:
            self._open_panel_animated()

    def _on_escape(self):
        if self._stack.currentIndex() == 1:
            self._show_chat_page()
        else:
            self._close_panel_animated()

    # ── MESSAGING ─────────────────────────────────────────────────────
    def _add_row(self, widget: QWidget, animate: bool = True):
        if self._empty_state.isVisible():
            self._empty_state.hide()

        row = _SlideInRow(widget)
        self._msg_layout.insertWidget(self._msg_layout.count() - 1, row)

        if animate:
            fx = QGraphicsOpacityEffect(row)
            fx.setOpacity(0.0)
            row.setGraphicsEffect(fx)

            fade = QPropertyAnimation(fx, b"opacity")
            fade.setStartValue(0.0); fade.setEndValue(1.0)
            fade.setDuration(180); fade.setEasingCurve(QEasingCurve.OutCubic)

            slide = QPropertyAnimation(row, b"offset")
            slide.setStartValue(12); slide.setEndValue(0)
            slide.setDuration(180); slide.setEasingCurve(QEasingCurve.OutCubic)

            group = QParallelAnimationGroup(row)
            group.addAnimation(fade)
            group.addAnimation(slide)

            def _detach_fx():
                try:
                    row.setGraphicsEffect(None)
                except RuntimeError:
                    pass

            group.finished.connect(_detach_fx)
            group.start()
            row._entrance_anim = group
        else:
            row._set_offset(0)

        row.show()
        self._scroll_to_bottom()
        QTimer.singleShot(50, self._scroll_to_bottom)

    def _add_message(self, role: str, text: str, animate: bool = True):
        self._add_row(MessageBubble(role, text), animate)

    def _add_status(self, text: str, persist: bool = True, animate: bool = True):
        self._add_row(StatusLine(text), animate)
        if persist:
            db_append("status", text)

    def _on_tool_status(self, text: str):
        self._add_status(text)

    def _on_reminder_requested(self, ms: int, message: str):
        # QTimer must live on the GUI thread — this slot always runs there.
        # message format: "<db_id>|<text>" (the id lets us mark it fired).
        rid = None
        if "|" in message:
            head, _, rest = message.partition("|")
            if head.isdigit():
                rid, message = int(head), rest
        print(f"[Kogane] reminder scheduled in {ms} ms: {message}", flush=True)
        QTimer.singleShot(ms, lambda: self._fire_reminder(message, rid))

    def _reschedule_saved_reminders(self):
        now = datetime.now()
        missed, scheduled = [], 0
        for rid, fire_at, msg in db_pending_reminders():
            if fire_at <= now:
                missed.append((rid, msg))
            else:
                ms = int((fire_at - now).total_seconds() * 1000)
                self._on_reminder_requested(ms, f"{rid}|{msg}")
                scheduled += 1
        if scheduled:
            print(f"[Kogane] rescheduled {scheduled} saved reminder(s)", flush=True)
        for rid, msg in missed:
            db_mark_reminder_fired(rid)
        if missed:
            names = "; ".join(m for _, m in missed[:5])
            QTimer.singleShot(2500, lambda: self._fire_reminder(
                f"Missed while offline: {names}", None))

    def _fire_reminder(self, message: str, rid: int | None = None):
        if Notification is not None:
            try:
                toast = Notification(app_id="Kogane",
                                      title="Kogane — Announcement",
                                      msg=message,
                                      icon=ICON_PATH)
                toast.set_audio(audio.Reminder, loop=False)
                toast.show()
            except Exception as _exc:
                print(f"[Kogane] toast failed: {_exc}", flush=True)
        if rid is not None:
            db_mark_reminder_fired(rid)
        if not self.face_widget.isVisible():
            self.face_widget.summon()
        else:
            self.face_widget.summon()  # replay the star burst as the alert
        # bounce() no-ops unless idle, so wait out the summon animation
        # (state == "summoning") before queuing the attention hop.
        QTimer.singleShot(
            int(kogane_animation.SUMMON_DURATION * 1000) + 30,
            self.face_widget.bounce)
        if not self._chat_open:
            self._open_panel_animated()
        self._add_message("ai", f"Announcement: {message}")
        db_append("assistant", f"Announcement: {message}")

    def _scroll_to_bottom(self):
        sb = self._scroll.verticalScrollBar()
        sb.setValue(sb.maximum())

    def _fill_input(self, text: str):
        self._input.setPlainText(text)
        self._input.setFocus()
        self._input.moveCursor(QTextCursor.End)

    def _set_thinking(self, on: bool):
        self._locked = on
        self._send_btn.setEnabled(not on)
        self._send_btn.set_thinking(on)
        if on:
            self._thinking_wrap.show()
            self._typing_bubble.start()
            kogane_animation.BOB_SPEED = self._THINKING_BOB_SPEED
        else:
            self._typing_bubble.stop()
            self._thinking_wrap.hide()
            kogane_animation.BOB_SPEED = self._DEFAULT_BOB_SPEED
        QTimer.singleShot(30, self._scroll_to_bottom)

    # ── ATTACHMENTS ───────────────────────────────────────────────────
    def _pick_attachment(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "Attach a file", "",
            "Supported files (*.txt *.md *.py *.js *.html *.css *.json *.csv "
            "*.log *.pdf *.png *.jpg *.jpeg *.webp);;All files (*)")
        if not path:
            return
        att = read_attachment(path)
        if att["kind"] == "error":
            self._add_status(att["content"], persist=False)
            return
        self._attachment = att
        self._attach_chip.setText(f"{att['name']}  ✕")
        self._attach_row.show()

    def _clear_attachment(self):
        self._attachment = None
        self._attach_row.hide()

    # ── SEND FLOW ─────────────────────────────────────────────────────
    def _on_send(self):
        text = self._input.toPlainText().strip()
        if self._locked:
            return
        if not text and self._attachment is None:
            return
        if not text:
            text = f"Look at the attached file."
        if len(text) > self._MAX_INPUT:
            text = text[:self._MAX_INPUT]
        self._input.clear()
        self._dispatch(text)

    def _on_send_btn_pressed(self):
        if self._send_btn_base_geo is None:
            self._send_btn_base_geo = self._send_btn.geometry()
        self._animate_send_btn_scale(0.92)

    def _on_send_btn_released(self):
        self._animate_send_btn_scale(1.0)

    def _animate_send_btn_scale(self, factor: float):
        base = self._send_btn_base_geo or self._send_btn.geometry()
        w = base.width() * factor
        h = base.height() * factor
        x = base.x() + (base.width() - w) / 2
        y = base.y() + (base.height() - h) / 2
        target = QRect(round(x), round(y), round(w), round(h))
        anim = QPropertyAnimation(self._send_btn, b"geometry")
        anim.setDuration(90)
        anim.setEasingCurve(QEasingCurve.OutCubic)
        anim.setEndValue(target)
        anim.start()
        self._send_scale_anim = anim

    def _dispatch(self, text: str):
        global _active_conv_id
        if _active_conv_id is None:
            _active_conv_id = db_new_conversation(_make_title(text))
        db_append("user", text)
        self._add_message("user", text)
        self._set_thinking(True)

        attachment = self._attachment
        self._clear_attachment()

        if attachment is not None and attachment["kind"] == "image":
            self._add_status(f"attached {attachment['name']}")
            b64 = attachment["b64"]
            fn = lambda status_cb, reminder_cb: ask_groq_with_image(text, b64)
        elif attachment is not None and attachment["kind"] == "text":
            self._add_status(f"attached {attachment['name']}")
            model_content = (
                f"{text}\n\n[Attached file: {attachment['name']}]\n"
                f"```\n{attachment['content']}\n```"
            )
            fn = lambda status_cb, reminder_cb: ask_groq(
                text, model_content, status_cb, reminder_cb)
        elif self._last_vision_b64:
            b64 = self._last_vision_b64
            self._last_vision_b64 = None
            fn = lambda status_cb, reminder_cb: ask_groq_with_image(text, b64)
        else:
            fn = lambda status_cb, reminder_cb: ask_groq(
                text, text, status_cb, reminder_cb)

        self._run_async(fn, self._on_reply, self._on_error)

    def _on_reply(self, reply: str):
        if reply == _SNIP_SENTINEL:
            # keep locked — the snip flow will unlock on completion/cancel
            self._start_snip(from_tool=True)
            return
        self._set_thinking(False)
        db_append("assistant", reply)
        self._add_message("ai", reply)
        # Deferred to the next event-loop turn: starting a new QThread here
        # synchronously, while still inside this worker thread's own
        # response_ready -> thread.quit()/deleteLater() signal chain,
        # crashed the interpreter (verified: the exact same call works fine
        # called directly, only crashes when nested in this callback).
        conv_id, reply_text = _active_conv_id, reply
        QTimer.singleShot(0, lambda: self._maybe_fire_title(conv_id, reply_text))

    def _maybe_fire_title(self, conv_id: int | None, assistant_reply: str):
        # Only the very first exchange earns a title call — never spend
        # one on a conversation that only has a single message so far.
        if conv_id is None or db_message_count(conv_id) != 2:
            return
        user_msgs = [m["content"] for m in db_load_messages(conv_id)
                     if m["role"] == "user"]
        if not user_msgs:
            return
        self._fire_title_request(conv_id, user_msgs[0], assistant_reply)

    def _fire_title_request(self, conv_id: int, user_msg: str, assistant_msg: str):
        """One-shot background title generation. Deliberately does NOT use
        self._thread/self._worker (those track the foreground chat reply) —
        a title call must never be cancelled by, or block, the next chat
        turn. Falls back to the existing truncated-prompt title on failure."""
        def fn(status_cb, reminder_cb):
            prompt = ("Write a 3-5 word title for this exchange. Title case. "
                      "No quotes, no punctuation. Reply with ONLY the title.")
            msgs = [
                {"role": "user", "content": user_msg},
                {"role": "assistant", "content": assistant_msg},
                {"role": "user", "content": prompt},
            ]
            r = groq_client.chat.completions.create(
                model=CHAT_MODEL, messages=msgs, max_tokens=200,
                timeout=GROQ_TIMEOUT_S)
            # CHAT_MODEL (gpt-oss) is a reasoning model — a small max_tokens
            # budget can be entirely consumed by hidden chain-of-thought,
            # leaving no room for the visible title. Give it headroom, and
            # strip any chain-of-thought that leaks into visible content.
            return strip_leaked_reasoning((r.choices[0].message.content or "").strip())

        thread = QThread(self)
        worker = GroqWorker(fn)
        worker.moveToThread(thread)
        thread.started.connect(worker.run)

        job = (thread, worker)
        self._title_jobs.append(job)

        def _cleanup_job():
            if job in self._title_jobs:
                self._title_jobs.remove(job)

        def on_title(raw: str):
            title = raw.strip().strip("\"'").replace("\n", " ").strip()
            title = title[:40].rstrip()
            if title:
                db_update_conversation_title(conv_id, title)
                print(f"[Kogane] AI title set for conversation {conv_id}: {title!r}",
                      flush=True)
            else:
                print(f"[Kogane] title generation returned empty — keeping prompt title",
                      flush=True)
            _cleanup_job()

        def on_title_error(msg: str):
            print(f"[Kogane] title generation failed, keeping prompt title: {msg}",
                  flush=True)
            _cleanup_job()

        worker.response_ready.connect(on_title)
        worker.error.connect(on_title_error)
        worker.response_ready.connect(thread.quit)
        worker.error.connect(thread.quit)
        worker.response_ready.connect(worker.deleteLater)
        worker.error.connect(worker.deleteLater)
        thread.finished.connect(thread.deleteLater)
        print(f"[Kogane] firing background title request for conversation {conv_id}",
              flush=True)
        thread.start()

    def _on_error(self, msg: str):
        self._set_thinking(False)
        self._add_message("ai", msg)
        self.face_widget.shake()

    def _clear_chat_view(self):
        # Layout: [0]=empty_state, [1]=stretch, [2..N-2]=rows, [N-1]=thinking
        while self._msg_layout.count() > 3:
            item = self._msg_layout.takeAt(2)
            if item.widget():
                item.widget().deleteLater()
        self._empty_state.show()

    def _new_conversation(self):
        global _active_conv_id
        _active_conv_id = None
        conversation_history.clear()
        self._last_vision_b64 = None
        self._clear_attachment()
        self._clear_chat_view()
        self._show_chat_page()

    # ── SCREEN SNIP ───────────────────────────────────────────────────
    def _start_snip(self, from_tool: bool = False):
        if self._locked and not from_tool:
            return
        if not self._chat_open:
            self._open_panel_animated()
        self.hide()
        self.face_widget.hide()  # keep Kogane out of her own screenshot
        QTimer.singleShot(200, self._selector.activate)

    def _on_snip_cancelled(self):
        self._set_thinking(False)
        self.show()
        self.face_widget.show()

    def _on_region_selected(self, x1, y1, x2, y2):
        self.show()
        self.face_widget.show()
        self._set_thinking(True)
        self._add_message("user", "screen capture")

        def capture(status_cb, reminder_cb):
            b64 = screen_to_b64(region=(x1, y1, x2, y2))
            return ask_groq_vision(b64)

        self._run_async(capture, self._on_snip_result, self._on_error)

    def _on_snip_result(self, reply: str):
        global _active_conv_id
        if _active_conv_id is None:
            _active_conv_id = db_new_conversation("Screen reading")
        db_append("assistant", reply)
        self._last_vision_b64 = None
        self._set_thinking(False)
        self._add_message("ai", reply)

    # ── VISIBILITY / QUIT ─────────────────────────────────────────────
    def _toggle_visibility(self):
        if self._chat_open:
            self._close_panel_animated()
        else:
            self._open_panel_animated()

    def _do_quit(self):
        self.face_widget.dismiss()
        QTimer.singleShot(400, QApplication.quit)

    # ── EVENT FILTER (Enter to send) ──────────────────────────────────
    def eventFilter(self, obj, event):
        from PyQt5.QtCore import QEvent
        if obj is self._input and event.type() == QEvent.KeyPress:
            if (event.key() in (Qt.Key_Return, Qt.Key_Enter)
                    and not (event.modifiers() & Qt.ShiftModifier)):
                self._on_send()
                return True
        return super().eventFilter(obj, event)

    def paintEvent(self, _):
        pass


# ── SYSTEM TRAY ───────────────────────────────────────────────────────
def _setup_tray(window: KoganeWindow):
    if pystray is None:
        print("[Kogane] Skipping tray (pystray unavailable)", flush=True)
        return None
    try:
        img = PILImage.open(ICON_PATH).convert("RGBA")
    except Exception:
        img = PILImage.new("RGBA", (64, 64), (232, 168, 37, 255))

    menu = pystray.Menu(
        pystray.MenuItem("Show / Hide",      lambda: window._toggle_sig.emit()),
        pystray.MenuItem("Screen snip",      lambda: window._trigger_snip_sig.emit()),
        pystray.MenuItem("New conversation", lambda: window._new_conv_sig.emit()),
        pystray.MenuItem("History",          lambda: window._history_sig.emit()),
        pystray.MenuItem("Quit",             lambda: window._quit_sig.emit()),
    )
    icon = pystray.Icon("Kogane", img, "Kogane", menu)
    threading.Thread(target=icon.run, daemon=True).start()
    return icon


# ── HOTKEYS ───────────────────────────────────────────────────────────
# Preferred: Windows RegisterHotKey via a native event filter. The OS only
# wakes Kogane when the exact combo fires — NO keyboard hook, no
# per-keystroke Python overhead, no input lag in games. The `keyboard`
# library hook is only a fallback if native registration fails.

_WM_HOTKEY   = 0x0312
_MOD_CONTROL = 0x0002
_MOD_SHIFT   = 0x0004
_MOD_NOREPEAT = 0x4000
_VK_SPACE    = 0x20
_VK_S        = 0x53

def _try_native_hotkeys(window: "KoganeWindow", app: QApplication) -> bool:
    if os.name != "nt":
        return False
    try:
        import ctypes
        import ctypes.wintypes
        from PyQt5.QtCore import QAbstractNativeEventFilter

        u32 = ctypes.windll.user32
        ok1 = u32.RegisterHotKey(None, 1, _MOD_CONTROL | _MOD_NOREPEAT, _VK_SPACE)
        ok2 = u32.RegisterHotKey(None, 2, _MOD_CONTROL | _MOD_SHIFT | _MOD_NOREPEAT, _VK_S)
        if not (ok1 and ok2):
            u32.UnregisterHotKey(None, 1)
            u32.UnregisterHotKey(None, 2)
            print("[Kogane] RegisterHotKey refused (combo in use?) — "
                  "falling back to keyboard hook", flush=True)
            return False

        class _HotkeyFilter(QAbstractNativeEventFilter):
            def nativeEventFilter(self, event_type, message):
                if event_type == b"windows_generic_MSG":
                    msg = ctypes.wintypes.MSG.from_address(int(message))
                    if msg.message == _WM_HOTKEY:
                        if msg.wParam == 1:
                            window._toggle_sig.emit()
                        elif msg.wParam == 2:
                            window._trigger_snip_sig.emit()
                return False, 0

        window._native_hotkey_filter = _HotkeyFilter()   # keep a reference
        app.installNativeEventFilter(window._native_hotkey_filter)
        print("[Kogane] Native hotkeys registered (RegisterHotKey — "
              "no keyboard hook, game-safe)", flush=True)
        return True
    except Exception as exc:
        print(f"[Kogane] native hotkeys failed: {exc}", flush=True)
        return False

def _hotkey_thread(window: KoganeWindow):
    if keyboard is None:
        print("[Kogane] Skipping global hotkeys (keyboard unavailable)", flush=True)
        return
    try:
        keyboard.add_hotkey(HOTKEY,      lambda: window._toggle_sig.emit())
        keyboard.add_hotkey(HOTKEY_SNIP, lambda: window._trigger_snip_sig.emit())
        print("[Kogane] WARNING: using keyboard-hook fallback for hotkeys — "
              "may cause input lag in games", flush=True)
    except Exception as exc:
        print(f"[Kogane] hotkey registration failed: {exc}", flush=True)
        return
    keyboard.wait()


# ── KEY SETUP ─────────────────────────────────────────────────────────
def _show_key_dialog() -> str:
    try:
        import tkinter as tk
        from tkinter import simpledialog
        root = tk.Tk(); root.withdraw()
        root.attributes("-topmost", True)
        key = simpledialog.askstring(
            "Kogane — Setup",
            "GROQ_API_KEY not set.\n\nPaste your Groq API key:",
            parent=root,
        )
        root.destroy()
        return (key or "").strip()
    except Exception as exc:
        print(f"[Kogane] key dialog error: {exc}", flush=True)
        return ""

def _save_key_to_env(key: str):
    env_path = os.path.join(BASE_DIR, ".env")
    lines = []
    if os.path.exists(env_path):
        with open(env_path) as f:
            lines = f.readlines()
    lines = [l for l in lines if not l.startswith("GROQ_API_KEY=")]
    lines.insert(0, f"GROQ_API_KEY={key}\n")
    with open(env_path, "w") as f:
        f.writelines(lines)


# ── ENTRY POINT ───────────────────────────────────────────────────────
def main():
    global groq_client, FONT_FAMILY

    api_key = GROQ_API_KEY
    if not api_key:
        api_key = _show_key_dialog()
        if not api_key:
            print("[Kogane] No API key — exiting", flush=True)
            sys.exit(1)
        _save_key_to_env(api_key)

    groq_client = Groq(api_key=api_key)

    QApplication.setAttribute(Qt.AA_EnableHighDpiScaling)
    QApplication.setAttribute(Qt.AA_UseHighDpiPixmaps)

    app = QApplication(sys.argv)
    app.setQuitOnLastWindowClosed(False)

    FONT_FAMILY = _load_font_family()
    app.setFont(_px_font(FS_BODY))
    print(f"[Kogane] QApplication.font() family = {app.font().family()!r}", flush=True)

    window = KoganeWindow()
    _setup_tray(window)
    if not _try_native_hotkeys(window, app):
        threading.Thread(target=_hotkey_thread, args=(window,), daemon=True).start()

    sys.exit(app.exec_())


if __name__ == "__main__":
    try:
        main()
    except Exception:
        print("\n[Kogane] FATAL:\n", flush=True)
        traceback.print_exc()
        try: input("Press Enter to exit…")
        except Exception: pass
        sys.exit(1)
